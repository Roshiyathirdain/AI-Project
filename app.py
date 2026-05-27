# app.py
import io
import os
import json
import time
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from mcp_engine import get_all_mcp_tools_sync, mcp_answer, _zai_chat
from utils import env

# ─── Page Configuration ────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Glimpse AI",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=DM+Sans:wght@400;500;700&display=swap');

    html, body, .stApp {
        background: #ffffff !important;
        color: #212121 !important;
        font-family: 'Inter', sans-serif;
    }

    /* ── Card Panel ── */
    .bi-panel {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 24px;
        margin-bottom: 20px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: border-color 0.2s ease, box-shadow 0.2s ease;
    }
    .bi-panel:hover { border-color: #d1d5db; box-shadow: 0 4px 12px rgba(0,0,0,0.08); }

    /* ── KPI Tile ── */
    .kpi-tile {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 16px 20px;
        text-align: center;
        position: relative;
        overflow: hidden;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    .kpi-tile:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0,0,0,0.08);
    }
    .kpi-tile::before {
        content: "";
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 4px;
    }
    .kpi-tile.blue::before  { background: #2563eb; }
    .kpi-tile.green::before { background: #10b981; }
    .kpi-tile.amber::before { background: #f59e0b; }
    .kpi-tile.red::before   { background: #ef4444; }
    .kpi-tile.purple::before{ background: #8b5cf6; }
    .kpi-tile.cyan::before  { background: #06b6d4; }
    
    .kpi-label {
        font-size: 0.75rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
        color: #6b7280;
        font-weight: 600;
        margin-bottom: 4px;
    }
    .kpi-value {
        font-size: 1.75rem;
        font-weight: 700;
        color: #111827;
        font-family: 'DM Sans', sans-serif;
        line-height: 1.1;
    }
    .kpi-sub {
        font-size: 0.7rem;
        color: #9ca3af;
        margin-top: 4px;
    }
    .kpi-sub .pos { color: #059669; }
    .kpi-sub .neg { color: #dc2626; }

    /* ── Headers ── */
    .chart-header {
        font-size: 0.75rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        color: #9ca3af;
        margin-bottom: 16px;
        display: flex;
        align-items: center;
        gap: 8px;
    }

    /* ── Record Card ── */
    .entity-card {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 16px;
        margin-bottom: 12px;
        border-left: 4px solid #2563eb;
        transition: background 0.2s ease, border-color 0.2s ease;
    }
    .entity-card:hover {
        background: #f9fafb;
    }
    .entity-key { color: #6b7280; font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.02em; }
    .entity-val { color: #111827; font-size: 0.88rem; font-weight: 500; }

    /* ── Sidebar ── */
    section[data-testid="stSidebar"] {
        background: #f7f7f8 !important;
        border-right: 1px solid #e5e7eb;
    }
    
    /* ── Chat Messages ── */
    .stChatMessage {
        background: transparent !important;
        border: none !important;
        padding: 24px 0 !important;
    }
    
    /* Assistant Background */
    div[data-testid="stChatMessage"] {
         border-bottom: 1px solid #f0f0f0;
    }

    .stMarkdown p { color: #374151; line-height: 1.75; font-size: 0.95rem; }

    /* ── Streamlit Overrides ── */
    .stAlert { border-radius: 12px !important; border: 1px solid #e5e7eb !important; background: #ffffff !important; }
    .stButton button { border-radius: 8px !important; }
    
    #MainMenu, footer { visibility: hidden; }
    .viewerBadge_container__r5tak { display: none !important; }
    
    header[data-testid="stHeader"] {
        background: #ffffff !important;
        border-bottom: 1px solid #e5e7eb;
    }

    .stTabs [data-baseweb="tab"] {
        font-size: 0.8rem !important;
        font-weight: 600 !important;
        color: #6b7280 !important;
    }
    .stTabs [aria-selected="true"] {
        color: #111827 !important;
        border-bottom-color: #2563eb !important;
    }
</style>
""", unsafe_allow_html=True)

# ─── Session State ─────────────────────────────────────────────────────────────
if "messages" not in st.session_state:
    st.session_state.messages = []
if "mcp_tools" not in st.session_state:
    st.session_state.mcp_tools = get_all_mcp_tools_sync()
if "session_files" not in st.session_state:
    st.session_state.session_files = {}   # {filename: text_content}
if "session_dfs" not in st.session_state:
    st.session_state.session_dfs = {}     # {filename: pd.DataFrame}
if "session_id" not in st.session_state:
    st.session_state.session_id = str(int(time.time()))

# ─── Persistent Chat History Logic ─────────────────────────────────────────────
HISTORY_FILE = "chat_history.json"

def save_chat_session():
    if not st.session_state.messages: return
    try:
        history = get_chat_history()
        # Create a title from the first user message
        first_user_msg = next((m["content"] for m in st.session_state.messages if m["role"] == "user"), "New Chat")
        title = first_user_msg[:40] + ("..." if len(first_user_msg) > 40 else "")
        
        current_sid = st.session_state.session_id
        
        # Check if session exists
        found = False
        for entry in history:
            if entry["id"] == current_sid:
                entry["title"] = title
                entry["messages"] = st.session_state.messages
                found = True
                break
        
        if not found:
            history.insert(0, {
                "id": current_sid, 
                "title": title, 
                "messages": st.session_state.messages,
                "dt": time.ctime()
            })
        
        # Keep only last 25 sessions
        history = history[:25]
        with open(HISTORY_FILE, "w") as f:
            json.dump(history, f)
    except: pass

def get_chat_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r") as f:
                return json.load(f)
        except: return []
    return []

def load_session(session_id):
    history = get_chat_history()
    for entry in history:
        if entry["id"] == session_id:
            st.session_state.messages = entry["messages"]
            st.session_state.session_id = entry["id"]
            return True
    return False

# ─── In-Memory File Text Extraction ────────────────────────────────────────────
FILE_ICONS = {"csv":"📊","xlsx":"📗","xls":"📗","pdf":"📕",
              "docx":"📘","doc":"📘","txt":"📄","md":"📝","json":"🗄️"}

def extract_text(uploaded_file) -> str:
    """Extract plain text from an uploaded file entirely in-memory."""
    name = uploaded_file.name
    ext  = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    raw  = uploaded_file.read()  # bytes

    try:
        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages), None

        elif ext in ("xlsx", "xls"):
            sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
            out = []
            all_dfs = []
            for sheet, df in sheets.items():
                # SAAS PRODUCTION GRADE: Robust Header Detection
                if any("Unnamed" in str(c) for c in df.columns) or df.columns.tolist() == list(range(len(df.columns))):
                    for i in range(min(5, len(df))):
                        row = df.iloc[i]
                        if row.count() >= 2:
                            df.columns = [str(x).strip() for x in row]
                            df = df.iloc[i+1:].reset_index(drop=True)
                            break
                out.append(f"=== Sheet: {sheet} ===\n{df.to_string(index=False)}")
                all_dfs.append(df)
            combined_df = pd.concat(all_dfs, ignore_index=True) if all_dfs else pd.DataFrame()
            return out, combined_df

        elif ext == "csv":
            df = pd.read_csv(io.BytesIO(raw))
            # Similar header logic for CSV
            if any("Unnamed" in str(c) for c in df.columns):
                potential_headers = df.iloc[0]
                if any(isinstance(x, str) and len(x) > 2 for x in potential_headers):
                    df.columns = df.iloc[0]
                    df = df.iloc[1:].reset_index(drop=True)
            return [df.to_string(index=False)], df

        elif ext in ("docx", "doc"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(raw))
                out = []
                
                # 1. Headers & Footers (Often contains Chapter Titles)
                for section in doc.sections:
                    for header in [section.header, section.first_page_header, section.even_page_header]:
                        if header: out.extend([p.text for p in header.paragraphs if p.text.strip()])
                    for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                        if footer: out.extend([p.text for p in footer.paragraphs if p.text.strip()])

                # 2. Main Body Paragraphs
                out.extend([p.text for p in doc.paragraphs if p.text.strip()])

                # 3. Tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text: out.append(row_text)
                
                # 4. Shapes/Textboxes (Often used for diagrams/notes)
                # Attempt to find text boxes in shapes
                try:
                    for shape in doc.inline_shapes:
                        pass # inline shapes usually don't have text we can jump to easily, but we try
                except: pass

                return "\n\n".join(out), None
            except Exception:
                # .doc (old binary format) — extract printable ASCII as fallback
                import re as _re
                decoded = raw.decode("latin-1", errors="ignore")
                # Extract runs of printable characters (length >= 4)
                words = _re.findall(r'[\x20-\x7E]{4,}', decoded)
                clean = " ".join(words)
                if len(clean) > 100:
                    return clean
                return (
                    "[Warning] Could not fully read this .doc file. "
                    "Please save it as .docx (Word 2007+) and re-upload for best results.\n\n"
                    + clean
                )

        else:  # txt, md, json, etc.
            return raw.decode("utf-8", errors="ignore"), None

    except Exception as e:
        return f"[Error reading {name}: {e}]", None


# ─── In-Memory File Q&A ─────────────────────────────────────────────────────────
def answer_from_session_files(question: str, session_files: dict, llm_model: str) -> dict:
    """Answer a question from in-memory uploaded files. Never reads from disk."""

    if not session_files:
        return {
            "answer": (
                "📁 **No files uploaded yet.**\n\n"
                "Please upload a PDF, Word, Excel, or CSV file using the **Upload Files** panel in the sidebar."
            ),
            "data": None, "sql": "", "visual_hint": {"show": False}
        }

    q_lower = question.lower()
    file_names = list(session_files.keys())

    # ── Pick the most relevant file (or combine all) ───────────────────────────
    target_name    = None
    target_content = ""

    for fname, content in session_files.items():
        stem = fname.rsplit(".", 1)[0].lower() if "." in fname else fname.lower()
        if fname.lower() in q_lower or stem in q_lower:
            target_name    = fname
            target_content = content
            break

    if not target_name:
        combined = [f"=== {fn} ===\n{ct}" for fn, ct in session_files.items()]
        target_content = "\n\n".join(combined)
        target_name    = ", ".join(file_names)

    # ── Detect meta/overview questions → use document start ───────────────────
    meta_patterns = [
        "about what", "discuss", "about this", "what is this", "what does this",
        "tell me about", "summarize", "summarise", "summary", "overview",
        "introduction", "topic", "subject", "this document", "the document",
        "this report", "the report", "what is it about", "explain this",
        "describe this", "document about", "report about",
    ]
    is_meta = any(p in q_lower for p in meta_patterns)

    # ── Build context — always give the LLM something to work with ─────────────
    paragraphs = [p.strip() for p in target_content.split("\n") if p.strip()]

    if is_meta:
        # Overview: start of document
        context = "\n".join(paragraphs[:50])[:3500]
    else:
        # ── Keyword scoring with Sliding Window ──────────────────────────────────
        import re as regex_match
        
        # Define keywords from the question to use in scoring
        keywords = [kw for kw in q_lower.split() if len(kw) > 3]
        
        # Enhanced Chapter Regex (Matches Ch 6, Chapter VI, Section 6.1, etc)
        chapter_pat = r'(chapter|section|part|unit)\s*([0-9]+|[ivxlc]+|[a-z])'
        user_chap = regex_match.search(chapter_pat, q_lower)
        
        scored_indices = []
        for i, p in enumerate(paragraphs):
            score = 0
            p_low = p.lower()
            
            if user_chap and user_chap.group(0) in p_low:
                score += 100 # High priority for exact chapter match
            
            # Context-sensitive match: if we found "Chapter 6", prioritize numbers like "6"
            if user_chap and user_chap.group(2) in p_low:
                score += 5
                
            score += sum(2 for kw in keywords if kw in p_low)
            if score > 0:
                scored_indices.append((score, i))

        if scored_indices:
            # Sort by score, keep top 15 matches
            scored_indices.sort(key=lambda x: x[0], reverse=True)
            top_indices = [idx for score, idx in scored_indices[:15]]
            
            # Build context using Sliding Window (pull in nearby lines for each match)
            final_context_blocks = []
            seen_indices = set()
            
            for idx in top_indices:
                # Take 2 paragraphs before and 6 paragraphs after the match
                for window_idx in range(max(0, idx - 2), min(len(paragraphs), idx + 7)):
                    if window_idx not in seen_indices:
                        final_context_blocks.append(paragraphs[window_idx])
                        seen_indices.add(window_idx)
            
            context = "\n".join(final_context_blocks)[:3500]
        else:
            # Fallback: Just give the beginning
            context = "\n".join(paragraphs[:30])[:3500]

    # ── Structure for LLM ─────────────────────────────────────────────────────
    df_info = ""
    if target_name in st.session_state.session_dfs:
        df_obj = st.session_state.session_dfs[target_name]
        df_info = (
            f"DATAFRAME INFO:\n"
            f"- Filename: {target_name}\n"
            f"- Columns: {df_obj.columns.tolist()}\n"
            f"- First 5 rows:\n{df_obj.head(5).to_string()}\n"
        )

    prompt = (
        "SYSTEM: You are a Senior SaaS Data Scientist. Answer ONLY based on the context.\n"
        f"CONTEXT EXCERPT:\n{context}\n\n"
        f"DATA INFO:\n{df_info}\n\n"
        "TASK: Answer the question using ONLY the info above. If missing, say 'NOT_FOUND_IN_DOCUMENTS'.\n"
        "PRODUCTION RULES:\n"
        "1. Write ONLY the code in ONE ```python block.\n"
        "2. SET: 'result' (answer), 'insights' (list), 'fig' (Plotly fig), 'prediction' (string).\n"
        "3. **VISUALIZATION**: Create a 'fig' (Plotly) ONLY if the data involves trends, comparisons, or distributions where a chart adds value. Avoid charts for simple lookups.\n"
        "4. Use 'parse_duration(text)'. No external imports.\n\n"
        f"USER QUESTION: {question}\n"
        "PYTHON CODE:"
    )

    try:
        raw_answer = _zai_chat(prompt, llm_model, timeout=300)
        clean = raw_answer.strip()

        # ── Execution Logic ──
        if "```python" in clean and target_name in st.session_state.session_dfs:
            calc_error = None
            code = ""
            try:
                import re as _re
                import pandas as pd
                import numpy as _np
                import plotly.express as px
                from datetime import datetime as _dt, timedelta as _td
                try:
                    from sklearn.linear_model import LinearRegression
                    from sklearn.cluster import KMeans
                except: pass
                
                # 1. Extraction
                code_match = _re.search(r"```python\n(.*?)```", clean, _re.DOTALL)
                if code_match:
                    code = code_match.group(1)
                    
                    # 2. Aggressive Sanitization
                    lines = code.split("\n")
                    clean_lines = []
                    for line in lines:
                        l_strip = line.strip()
                        if l_strip.startswith(("import pandas", "import sklearn", "df = pd.read", "print(")):
                            continue
                        clean_lines.append(line)
                    code = "\n".join(clean_lines)
                    
                    # 3. Production Helper
                    def parse_duration(val):
                        try:
                            val = str(val).lower().replace(" ", "").replace(":", ".")
                            if "-" not in val: return 0.0
                            parts = val.split("-")
                            s = _dt.strptime(parts[0], "%I.%M%p")
                            e = _dt.strptime(parts[1], "%I.%M%p")
                            if e < s: e += _td(days=1)
                            return (e - s).total_seconds() / 3600.0
                        except: return 0.0

                    # 4. Mock Environment
                    df_live = st.session_state.session_dfs[target_name]
                    exec_scope = {
                        "df": df_live, "pd": pd, "np": _np, "px": px, "re": _re,
                        "datetime": _dt, "timedelta": _td,
                        "parse_duration": parse_duration, 
                        "result": None, "insights": [], "fig": None, "prediction": None
                    }
                    
                    # Try to add sklearn if possible
                    try:
                        from sklearn.linear_model import LinearRegression as LR
                        exec_scope["LinearRegression"] = LR
                    except: pass

                    exec(code, exec_scope)
                    
                    if exec_scope.get("result") is not None:
                        final_res = exec_scope.get("result")
                        # Grounding Check
                        if str(final_res).upper() == "NOT_FOUND_IN_DOCUMENTS":
                            return {
                                "answer": "❌ **Information Not Found in Documents**\n\nI couldn't find any information locally related to your question. Please try asking about something contained in your uploaded files.",
                                "data": None,
                                "sql": f"Checked: {target_name}",
                                "visual_hint": {"show": False}
                            }

                        # Prepare data for return (for persistence in chat history)
                        export_data = None
                        if isinstance(final_res, (pd.DataFrame, pd.Series, list)):
                            export_data = pd.DataFrame(final_res).to_dict(orient='records')
                        
                        ans_text = f"**Analysis Result:** {final_res if not isinstance(final_res, (pd.DataFrame, pd.Series, list)) else 'See analytical breakdown below'}"
                        
                        # Capture AI generated insights and predictions
                        insights = exec_scope.get("insights", [])
                        pred = exec_scope.get("prediction")
                        
                        return {
                            "answer": ans_text,
                            "data": export_data,
                            "insights": insights,
                            "prediction": pred,
                            "sql": f"SaaS Analysis Logic",
                            "visual_hint": {"show": True, "type": "Auto"} if exec_scope.get("fig") else {"show": False}
                        }
            except Exception as e:
                calc_error = str(e)
            
            if calc_error:
                with st.expander("🔍 Calculation Diagnostic", expanded=False):
                    st.error(f"Error: {calc_error}")
                    st.code(code)

        # Show the actual context use for debugging in the UI
        with st.expander("🔍 View Source Context (Paragraphs Used)", expanded=False):
            st.text(context)

        if "NOT_FOUND" in clean.upper() or len(clean) < 4:
            return {
                "answer": (
                    f"❌ **Not found in `{target_name}`.**\n\n"
                    "The uploaded document does not contain the answer to your question.\n\n"
                    "*Try asking something else, or upload a different file.*"
                ),
                "data": None,
                "sql": f"File: {target_name}",
                "visual_hint": {"show": False}
            }

        return {"answer": clean, "data": None, "sql": f"File: {target_name}", "visual_hint": {"show": False}}

    except Exception as e:
        # LLM failed → return raw content so user still gets something
        fallback = "\n".join(paragraphs[:12])
        return {
            "answer": (
                f"⚠️ *LLM unavailable ({e}). Showing raw document content:*\n\n"
                f"**{target_name}**\n\n{fallback}"
            ),
            "data": None,
            "sql": f"File: {target_name}",
            "visual_hint": {"show": False}
        }

# ─── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("🧩 BI Engine")
    st.caption("Powered by MCP & Z.ai Intelligence")
    
    # ── New Chat Button ──────────────────────────────────────────────
    if st.button("➕ New Chat", use_container_width=True):
        st.session_state.messages = []
        # Reset session ID for a truly new chat
        st.session_state.session_id = str(int(time.time()))
        st.rerun()

    st.divider()

    # ── Chat History List (Persistent Sessions) ──────────────────────────────
    st.markdown("### 🕑 Recent Chats")
    history = get_chat_history()
    
    for entry in history:
        # Highlight active session
        is_active = (entry['id'] == st.session_state.session_id)
        icon = "🗨️" if is_active else "💬"
        btn_label = f"{icon} {entry['title']}"
        
        if st.button(btn_label, key=f"hist_{entry['id']}", use_container_width=True):
            if load_session(entry['id']):
                st.rerun()
            
    if history and st.button("🗑️ Clear All History", use_container_width=True):
        if os.path.exists(HISTORY_FILE): os.remove(HISTORY_FILE)
        st.session_state.messages = []
        st.session_state.session_id = str(int(time.time()))
        st.rerun()

    st.divider()

    # ── Dynamic Model List ─────────────────────────────────────────────────────
    # ── Unified Model Selection ──────────────────────────────────────
    from mcp_engine import ZAI_API_KEY
    if ZAI_API_KEY:
        model_options = ["qwen3.5:cloud", "qwen2.5:latest", "llama3.2:1b", "glm-4.7", "glm-4.6", "glm-5", "glm-4.5"]
        st.info("🚀 **Hybrid Intelligence Active (Local + Cloud)**")
    else:
        model_options = ["qwen3.5:cloud", "qwen2.5:latest", "llama3.2:1b", "llama3.2:3b", "deepseek-r1:7b", "llama3.1:8b"]
    
    model_index = 0

    model = st.selectbox(
        "🤖 Intelligence Model",
        model_options,
        index=model_index
    )

    # Data Source selection
    data_source = st.selectbox(
        "📊 Data Source", 
        ["Auto", "Postgres", "MySQL", "SharePoint", "Local Files"],
        index=1, # Default to Postgres as requested
        help="Choose a specific system to query, or let the AI detect it automatically."
    )

    # ── File Upload Panel (only when Local Files is chosen) ────────────────────
    if data_source == "Local Files":
        st.divider()
        st.subheader("📎 Upload Files")
        st.caption("Files are read in-memory. Nothing is saved to disk.")

        uploaded = st.file_uploader(
            "Drop files here",
            type=["csv", "xlsx", "xls", "pdf", "docx", "doc", "txt", "md", "json"],
            accept_multiple_files=True,
            label_visibility="collapsed"
        )

        if uploaded:
            new_count = 0
            for f in uploaded:
                if f.name not in st.session_state.session_files:
                    text_data, df_data = extract_text(f)
                    # For sheets list, join if it's a list
                    st.session_state.session_files[f.name] = "\n".join(text_data) if isinstance(text_data, list) else text_data
                    if df_data is not None:
                        st.session_state.session_dfs[f.name] = df_data
                    new_count += 1
            if new_count:
                st.success(f"✅ {new_count} new file(s) loaded into memory")

        # Show currently loaded files
        sf = st.session_state.session_files
        if sf:
            st.markdown(f"**{len(sf)} file(s) loaded this session:**")
            for fname in sf:
                ext  = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
                icon = FILE_ICONS.get(ext, "📎")
                char_count = len(sf[fname])
                st.markdown(f"{icon} `{fname}` — *{char_count:,} chars*")

            if st.button("🗑️ Clear Uploaded Files"):
                st.session_state.session_files = {}
                st.session_state.session_dfs = {}
                st.rerun()
        else:
            st.info("No files loaded. Upload above to get started.")

    st.divider()
    st.subheader("📡 Distributed Core")
    if st.session_state.mcp_tools:
        st.success("✨ **Glimpse AI HTTP Gateway** (SSE/HTTP)")
        st.caption(f"Broadcasting {len(st.session_state.mcp_tools)} tools across company.")
        if st.button("🔄 Refresh Tools", width="stretch"):
            st.session_state.mcp_tools = get_all_mcp_tools_sync()
            st.rerun()
    else:
        st.error("❌ Gateway Offline")
        if st.button("📡 Reconnect to Gateway", width="stretch"):
            st.session_state.mcp_tools = get_all_mcp_tools_sync()
            st.rerun()

    st.divider()


# ─── Premium Header ────────────────────────────────────────────────────────────
top_c1, top_c2 = st.columns([0.7, 0.3])
with top_c1:
    st.title("✨ Glimpse AI")
    st.markdown("##### *Production SaaS Interface — Active Insight Engine*")
with top_c2:
    st.markdown("<br>", unsafe_allow_html=True)
    status_color = "#008a66" if st.session_state.mcp_tools else "#ef4444"
    status_bg = "rgba(0, 204, 150, 0.05)" if st.session_state.mcp_tools else "rgba(239, 68, 68, 0.05)"
    status_border = "rgba(0, 204, 150, 0.4)" if st.session_state.mcp_tools else "rgba(239, 68, 68, 0.4)"
    status_text = "🟢 AGENT ONLINE" if st.session_state.mcp_tools else "🔴 GATEWAY OFFLINE"
    conn_text = "MCP Gateway Connected" if st.session_state.mcp_tools else "Connection Lost"
    
    st.markdown(f"""
        <div style="background: {status_bg}; border: 1px solid {status_border}; border-radius: 12px; padding: 10px; text-align: center;">
            <span style="color: {status_color}; font-weight: 700;">{status_text}</span><br>
            <span style="color: rgba(0,0,0,0.5); font-size: 0.8rem;">{conn_text}</span>
        </div>
    """, unsafe_allow_html=True)

st.write("") # Spacer


# ─── Dashboard helpers ─────────────────────────────────────────────────────────
# ─── Question Intent Classification Engine ──────────────────────────────────
# This classifies every question into a visual mode WITHOUT calling the LLM.
# This is the production-grade approach used by tools like Tableau/PowerBI AI.

def classify_visual_intent(question: str, df) -> dict:
    """
    Returns a dict:
    {
      'mode': 'PLAIN' | 'TABLE' | 'KPI+TABLE' | 'CHART+KPI+TABLE',
      'chart_type': 'Bar' | 'Line' | 'Pie' | None,
      'reason': str   # dev debug string
    }
    
    PLAIN       → no data/table shown at all (pure text answer)
    TABLE       → just show the data table (listing, lookup, detail)
    KPI+TABLE   → financial/aggregate queries with totals
    CHART+KPI+TABLE → analytical/comparative/trends (full dashboard)
    """
    q = question.lower().strip()
    rows = len(df) if df is not None else 0
    cols = list(df.columns) if df is not None else []

    import numpy as np
    num_cols = [c for c in cols if df[c].dtype in [np.int64, np.float64, np.int32, np.float32]] if df is not None else []
    cat_cols = [c for c in cols if df[c].dtype == object] if df is not None else []

    # ── 1. PLAIN: conversation / capability questions (no data needed) ────────
    plain_kw = [
        "hello", "hi ", "hey ", "who are you", "what can you", "help me",
        "capabilities", "what data", "how do you", "what is", "tell me about",
        "explain", "define"
    ]
    if any(k in q for k in plain_kw):
        return {"mode": "PLAIN", "chart_type": None, "reason": "greeting/meta"}

    # ── 2. PLAIN: single-value aggregate (no visual needed) ───────────────────
    single_agg_kw = [
        "total payroll", "total salary", "salary total", "payroll total",
        "total revenue", "total count", "how many", "count of",
        "average salary", "avg salary", "max salary", "min salary",
        "highest paid", "lowest paid", "what is the total",
    ]
    if any(k in q for k in single_agg_kw):
        # If the result is only 1 row (single number), PLAIN is best
        if rows <= 1:
            return {"mode": "PLAIN", "chart_type": None, "reason": "single-value aggregate"}
        # If multiple rows returned (e.g. dept-wise), show CHART+KPI+TABLE
        return {"mode": "CHART+KPI+TABLE", "chart_type": "Bar", "reason": "multi-row aggregate"}

    # ── 3. CHART: explicit charting request ───────────────────────────────────
    chart_explicit_kw = [
        "chart", "graph", "plot", "visualize", "visualise",
        "show graph", "draw chart", "bar chart", "pie chart", "line chart"
    ]
    if any(k in q for k in chart_explicit_kw):
        chart_type = "Bar"
        if "line" in q or "trend" in q: chart_type = "Line"
        if "pie" in q or "donut" in q: chart_type = "Pie"
        return {"mode": "CHART+KPI+TABLE", "chart_type": chart_type, "reason": "explicit chart request"}

    # ── 4. CHART: comparison / trend / analysis ───────────────────────────────
    comparison_kw = [
        "profit by", "sales by", "revenue by", "expenses by", "cost by",
        "by region", "by department", "by category", "by product",
        "by month", "by year", "by quarter", "by date", "monthly", "quarterly", "yearly",
        "compare", "comparison", "vs ", "versus", "breakdown",
        "trend", "over time", "growth", "decline", "top 10", "top 5",
        "ranking", "ranked", "distribution", "share",
        "profit", "revenue", "analyze", "analyse", "analysis"
    ]
    if any(k in q for k in comparison_kw):
        chart_type = "Bar"
        if any(k in q for k in ["trend", "over time", "monthly", "quarterly", "yearly", "by month", "by year", "by date"]):
            chart_type = "Line"
        if any(k in q for k in ["share", "percent", "proportion", "breakdown", "distribution", "pie"]):
            chart_type = "Pie"
        return {"mode": "CHART+KPI+TABLE", "chart_type": chart_type, "reason": "comparison/analysis"}

    # ── 5. CHART: if data has both numeric+categorical cols and 2+ rows ───────
    if rows >= 2 and num_cols and cat_cols:
        # Financial/metric keywords
        financial_kw = ["sales", "revenue", "profit", "expense", "salary", "pay",
                        "budget", "cost", "price", "amount", "total", "sum"]
        if any(k in q for k in financial_kw):
            return {"mode": "CHART+KPI+TABLE", "chart_type": "Bar", "reason": "financial with structure"}

    # ── 6. KPI+TABLE: queries with numeric data (aggregate-ish) ──────────────
    kpi_kw = [
        "payroll", "salary", "wage", "income", "inventory", "stock",
        "invoice", "payment", "budget", "revenue", "sales", "profit",
        "expense", "cost", "price", "amount"
    ]
    if num_cols and any(k in q for k in kpi_kw):
        return {"mode": "KPI+TABLE", "chart_type": None, "reason": "numeric/financial data"}

    # ── 7. TABLE: listing / lookup (the most common query type) ──────────────
    # "show all employees", "list products", "get customers", etc.
    listing_kw = [
        "show", "list", "get", "display", "fetch", "view",
        "all employees", "all products", "all customers", "all users",
        "all orders", "all invoices", "all projects", "all departments",
        "all payroll", "all sales", "all inventory", "all suppliers",
        "employees", "members", "staff", "personnel",
    ]
    # ── 7. TABLE with Smart Visuals: listing with grouping potential ────────
    if any(k in q for k in listing_kw):
        if rows >= 3 and cat_cols:
             return {"mode": "CHART+KPI+TABLE", "chart_type": "Bar", "reason": "listing with categorization"}
        return {"mode": "TABLE", "chart_type": None, "reason": "pure listing"}

    # ── 8. If we have data with numeric cols, show KPI+TABLE ──────────────────
    if num_cols:
        return {"mode": "KPI+TABLE", "chart_type": None, "reason": "has metrics"}

    # ── 9. Default: just TABLE ────────────────────────────────────────────────
    return {"mode": "TABLE", "chart_type": None, "reason": "default"}


def get_meaningful_cols(df):
    """Filter to meaningful business columns only."""
    junk = ["id", "pk", "index", "phone", "mobile", "zip", "pin", "postal",
            "serial", "fax", "created_at", "updated_at"]
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    cat_cols = df.select_dtypes(include=["object"]).columns.tolist()
    clean_num = [c for c in num_cols if not any(x in c.lower() for x in junk)]
    priority_cat = ["name", "title", "role", "dept", "category", "status", "type"]
    clean_cat = sorted(
        [c for c in cat_cols if not any(x in c.lower() for x in junk)],
        key=lambda c: any(p in c.lower() for p in priority_cat),
        reverse=True
    )
    return clean_num, clean_cat


def _plotly_layout():
    return dict(
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(248,249,250,1)',
        font=dict(family="Inter", size=11, color="#4b5563"),
        title_font=dict(size=14, family="Inter", color="#111827"),
        margin=dict(t=40, b=30, l=10, r=10),
        xaxis=dict(gridcolor="#e5e7eb", linecolor="#e5e7eb", tickfont=dict(color="#6b7280")),
        yaxis=dict(gridcolor="#e5e7eb", linecolor="#e5e7eb", tickfont=dict(color="#6b7280")),
    )

BI_COLORS = ["#2563eb", "#10b981", "#f59e0b", "#8b5cf6", "#ef4444", "#06b6d4", "#ec4899"]
TILE_CLASSES = ["blue", "green", "amber", "purple", "red", "cyan"]


def render_kpi_tiles(df):
    """Power BI-style KPI metric tiles — only for numeric data."""
    num_cols, cat_cols = get_meaningful_cols(df)
    if not num_cols:
        # Fallback for listings: show record count as a KPI
        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
        c1, _ = st.columns([1, 4])
        with c1:
            st.markdown(f"""
                <div class="kpi-tile blue">
                    <div class="kpi-label">Total Records</div>
                    <div class="kpi-value">{len(df)}</div>
                    <div class="kpi-sub">entries analyzed</div>
                </div>
            """, unsafe_allow_html=True)
        st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)
        return
    st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
    n = min(len(num_cols), 5)
    kpi_cols = st.columns(n)
    for i, col in enumerate(num_cols[:n]):
        val = df[col].sum()
        avg = df[col].mean()
        tile_cls = TILE_CLASSES[i % len(TILE_CLASSES)]
        disp = f"{val:,.0f}" if val >= 1 else f"{val:.4f}"
        kpi_cols[i].markdown(f"""
            <div class="kpi-tile {tile_cls}">
                <div class="kpi-label">{col.replace('_',' ')}</div>
                <div class="kpi-value">{disp}</div>
                <div class="kpi-sub">avg {avg:,.1f}</div>
            </div>
        """, unsafe_allow_html=True)
    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)


def render_data_table(df):
    """Always-visible styled data table."""
    if df.empty:
        return
    num_cols, _ = get_meaningful_cols(df)
    try:
        styled = df.style.background_gradient(cmap="Blues", subset=num_cols[:2] if num_cols else [])
        st.dataframe(styled, use_container_width=True, height=min(400, 38 + len(df) * 36))
    except Exception:
        st.dataframe(df, use_container_width=True)


def render_chart(df, question, chart_type="Bar", visual_hint=None):
    """Render a single, intelligent chart."""
    if df.empty:
        return
    num_cols, cat_cols = get_meaningful_cols(df)
    
    # POWER BI LOGIC: Even if no numeric columns, we can show distribution (counts)
    is_frequency_chart = not num_cols and cat_cols
    if not num_cols and not cat_cols:
        return  # Purely textual data

    layout = _plotly_layout()

    # Resolve X axis (category)
    x_col = visual_hint.get("x") if visual_hint else None
    if not x_col or x_col not in df.columns:
        q_lower = question.lower()
        date_cols = [c for c in df.columns if any(x in c.lower() for x in ["date", "month", "year", "time"])]
        if chart_type == "Line" and date_cols:
            x_col = date_cols[0]
        elif cat_cols:
            x_col = cat_cols[0]
        else:
            x_col = df.columns[0]

    # Resolve Y axis (metric)
    y_col = visual_hint.get("y") if visual_hint else None
    if not y_col or y_col not in df.columns:
        if num_cols:
            y_col = num_cols[0]
        elif cat_cols:
            y_col = "Count"
            chart_df = df.groupby(x_col).size().reset_index(name="Count")
        else:
            y_col = None
    
    if not y_col:
        return

    # Chart title
    title_text = (visual_hint.get("title") if visual_hint else None) or \
                 f"{y_col.replace('_', ' ').upper()} BY {x_col.replace('_', ' ').upper()}"
    st.markdown(f"""<p class="chart-header">📊 &nbsp;{title_text}</p>""", unsafe_allow_html=True)

    if not 'chart_df' in locals():
        chart_df = df.copy()
    # Sort for better readability
    try:
        chart_df = chart_df.sort_values(y_col, ascending=False)
    except Exception:
        pass

    try:
        if chart_type == "Line":
            try: chart_df = chart_df.sort_values(x_col)
            except: pass
            fig = px.line(chart_df.head(30), x=x_col, y=y_col, markers=True,
                          color_discrete_sequence=["#2563eb"],
                          labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})
            fig.update_traces(line=dict(width=2.5), marker=dict(size=7))

        elif chart_type == "Pie":
            top_df = chart_df.nlargest(8, y_col)
            fig = px.pie(top_df, values=y_col, names=x_col, hole=0.45,
                         color_discrete_sequence=BI_COLORS)
            fig.update_traces(textposition='inside', textinfo='percent+label')

        else:  # Default: Bar
            # Production SaaS: Always limit to top 15 for readability, unless explicitly scrolled
            plot_df = chart_df.head(15) if len(chart_df) > 15 else chart_df
            
            # Use horizontal bars if too many items or long labels
            orientation = 'v'
            if len(plot_df) > 8 or plot_df[x_col].astype(str).str.len().max() > 12:
                orientation = 'h'
                # Swap X and Y for horizontal
                fig = px.bar(plot_df, x=y_col, y=x_col, text_auto='.3s',
                             color=y_col, orientation='h',
                             color_continuous_scale=[[0, "#1e3a6e"], [1, "#3b82f6"]],
                             labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})
                fig.update_layout(yaxis={'categoryorder':'total ascending'})
            else:
                fig = px.bar(plot_df, x=x_col, y=y_col, text_auto='.3s',
                             color=y_col,
                             color_continuous_scale=[[0, "#1e3a6e"], [1, "#3b82f6"]],
                             labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})

        fig.update_layout(**layout, coloraxis_showscale=False)
        st.plotly_chart(fig, use_container_width=True)

    except Exception as e:
        st.caption(f"Chart could not be rendered: {e}")


def render_response(df, question, model, visual_hint=None, cached_insights=None, cached_forecast=None):
    """
    Master render function — decides exactly what to show based on question intent.
    Returns (insights, forecast) for caching.
    """
    if df is None or df.empty:
        return None, None

    intent = classify_visual_intent(question, df)
    mode = intent["mode"]
    chart_type = intent["chart_type"]

    # Override with AI visual_hint if present and valid
    if visual_hint and isinstance(visual_hint, dict):
        if visual_hint.get("show"):
            mode = "CHART+KPI+TABLE"
            chart_type = visual_hint.get("type", "Bar")
        # Removed the 'False' downgrade to ensure PowerBI-style visual consistency

    insights = cached_insights
    forecast = cached_forecast

    # ── MODE: TABLE only (listings, lookups, directory queries) ──────────────
    if mode == "TABLE":
        num_cols, _ = get_meaningful_cols(df)
        # Only show record count tile — no KPIs, no charts
        st.markdown(f"""
            <div style="display:flex; gap:12px; margin-bottom:12px;">
                <div style="background:#f0f9ff; border:1px solid #bae6fd; border-radius:8px;
                            padding:8px 16px; font-size:0.82rem; color:#0369a1; font-weight:600;">
                    📋 {len(df)} records found
                </div>
                <div style="background:#f0fdf4; border:1px solid #bbf7d0; border-radius:8px;
                            padding:8px 16px; font-size:0.82rem; color:#166534; font-weight:600;">
                    🗂️ {len(df.columns)} columns
                </div>
            </div>
        """, unsafe_allow_html=True)
        render_data_table(df)
        return None, None

    # ── MODE: KPI + TABLE (financial, payroll, inventory) ────────────────────
    elif mode == "KPI+TABLE":
        render_kpi_tiles(df)
        render_data_table(df)
        return None, None

    # ── MODE: CHART + KPI + TABLE (comparative, trend, analytical) ───────────
    elif mode == "CHART+KPI+TABLE":
        # 1. KPI row
        render_kpi_tiles(df)
        # 2. Chart
        render_chart(df, question, chart_type, visual_hint)
        # 3. AI Insight Panel (only in full dashboard mode)
        if not insights or not forecast:
            sample = df.head(8).to_dict(orient='records')
            insight_prompt = f"""You are a BI analyst. Analyze this data.
Question: "{question}"
Data ({len(df)} rows): {json.dumps(sample)}

Write 2 bullet point insights and 1 forward-looking forecast.
Format EXACTLY as:
INSIGHTS:
- [insight 1]
- [insight 2]

FORECAST:
[one sentence prediction]
NO filler text."""
            try:
                raw = _zai_chat(insight_prompt, model, timeout=20)
                parts = raw.split("FORECAST:")
                insights = parts[0].replace("INSIGHTS:", "").strip()
                forecast = parts[1].strip() if len(parts) > 1 else ""
            except Exception:
                insights = f"• {len(df)} records analyzed from your query."
                forecast = "No forecast available — check Ollama connection."

        if insights:
            c1, c2 = st.columns([0.6, 0.4])
            with c1:
                st.markdown(f"""
                    <div style="background:rgba(37,99,235,0.04); border-left:4px solid #2563eb;
                                padding:14px; border-radius:8px; border:1px solid rgba(37,99,235,0.12);">
                        <div style="font-weight:700; color:#1e3a8a; margin-bottom:8px; font-size:0.78rem;
                                    text-transform:uppercase; letter-spacing:0.06em;">🧠 Agent Insights</div>
                        <div style="font-size:0.88rem; color:#374151; line-height:1.6;">{insights}</div>
                    </div>
                """, unsafe_allow_html=True)
            if forecast:
                with c2:
                    st.markdown(f"""
                        <div style="background:rgba(16,185,129,0.04); border-left:4px solid #10b981;
                                    padding:14px; border-radius:8px; border:1px solid rgba(16,185,129,0.12);">
                            <div style="font-weight:700; color:#065f46; margin-bottom:8px; font-size:0.78rem;
                                        text-transform:uppercase; letter-spacing:0.06em;">🔮 Forecast</div>
                            <div style="font-size:0.88rem; color:#065f46; line-height:1.6; font-style:italic;">{forecast}</div>
                        </div>
                    """, unsafe_allow_html=True)

        # 4. Data table (collapsed by default when chart is shown)
        with st.expander("🗄️ View Full Data Table", expanded=False):
            render_data_table(df)
        return insights, forecast

    # Fallback — show table
    render_data_table(df)
    return None, None


# ─── Chat Interface ─────────────────────────────────────────────────────────────
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        if m["role"] == "assistant":
            content = m['content']
            # Strip old debug prefix if present
            clean = content.replace("**ANSWER:** ", "").replace("ANSWER: ", "")
            st.markdown(clean)
        else:
            st.markdown(m["content"])
        sql_val = m.get("sql", "")
        if sql_val and sql_val not in ("General Conversation", "", None):
            is_tool = any(str(sql_val).startswith(p) for p in ("Tool:", "File:", "File "))
            label   = "🔧 Source" if is_tool else "🔍 SQL Query"
            with st.expander(label, expanded=False):
                st.code(sql_val, language="text" if is_tool else "sql")
        if m.get("data") and len(m["data"]) > 0:
            try:
                _df = pd.DataFrame(m["data"])
                orig_q = m.get("orig_question", m["content"])
                render_response(
                    _df, orig_q, model,
                    visual_hint=m.get("visual_hint"),
                    cached_insights=m.get("insights"),
                    cached_forecast=m.get("prediction")
                )
                st.download_button(
                    label="📥 Export CSV",
                    data=_df.to_csv(index=False).encode('utf-8'),
                    file_name=f"glimpse_export_{int(time.time())}.csv",
                    mime='text/csv',
                    key=f"dl_{m.get('id', hash(m['content']))}"
                )
            except Exception as _e:
                st.dataframe(pd.DataFrame(m["data"]), use_container_width=True)


prompt = st.chat_input("Ask a question about your data or uploaded files...")

if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        spinner_msg = (
            "📄 Searching uploaded files..."
            if data_source == "Local Files"
            else f"🔍 Analysing {data_source} data..."
            if data_source != "Auto"
            else "🤖 Consulting Intelligence..."
        )
        with st.spinner(spinner_msg):
            try:
                # ── Local Files: answer from in-memory content ─────────────
                if data_source == "Local Files":
                    res = answer_from_session_files(
                        prompt,
                        st.session_state.session_files,
                        model
                    )
                # ── All other sources: go through MCP engine ───────────────
                else:
                    res = mcp_answer(
                        prompt,
                        st.session_state.mcp_tools,
                        model,
                        data_source=data_source
                    )

                ans  = str(res.get("answer", "No answer returned.") or "No content found.")
                data = res.get("data", None)
                sql  = res.get("sql", "")
                v_hint = res.get("visual_hint", {})

                # ── 1. Primary Answer Text ────────────────────────────────────
                st.markdown(ans)

                # ── 2. SQL Source (collapsed) ─────────────────────────────────
                if sql and sql not in ("General Conversation", "", None):
                    is_tool = any(str(sql).startswith(p) for p in ("Tool:", "File:", "File "))
                    label   = "🔧 Source" if is_tool else "🔍 SQL Query"
                    with st.expander(label, expanded=False):
                        st.code(sql, language="text" if is_tool else "sql")

                # ── 3. Data Visualization (intent-driven) ─────────────────────
                saved_insights, saved_forecast = None, None
                if data and len(data) > 0:
                    df = pd.DataFrame(data)
                    try:
                        saved_insights, saved_forecast = render_response(
                            df, prompt, model,
                            visual_hint=v_hint
                        )
                        st.download_button(
                            label="📥 Export CSV",
                            data=df.to_csv(index=False).encode('utf-8'),
                            file_name=f"export_{int(time.time())}.csv",
                            mime='text/csv'
                        )
                    except Exception as e:
                        st.warning(f"⚠️ Visualization error: {e}")
                        st.dataframe(df, use_container_width=True)

                st.session_state.messages.append({
                    "role": "assistant",
                    "content": ans,
                    "data": data,
                    "orig_question": prompt,
                    "insights": saved_insights,
                    "prediction": saved_forecast,
                    "visual_hint": v_hint,
                    "sql": sql
                })

            except Exception as e:
                err_msg = f"❌ Analysis failed: {e}"
                st.error(err_msg)
                st.session_state.messages.append({"role": "assistant", "content": err_msg})
        # Save to history file for persistence
        save_chat_session()
        st.rerun()
