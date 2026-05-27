# app.py
import io
import os
import json
import time
import html
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from mcp_engine import get_all_mcp_tools_sync, mcp_answer, _gemini_chat
from utils import env

# ─── Page Configuration ────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Glimpse AI",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=DM+Sans:wght@400;500;700&display=swap');

    html, body, .stApp {
        background: #ffffff !important;
        color: #212121 !important;
        font-family: 'Inter', sans-serif;
    }

    /* ── Card Panel ── */
    .bi-panel {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 24px;
        margin-bottom: 20px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: border-color 0.2s ease, box-shadow 0.2s ease;
    }
    .bi-panel:hover { border-color: #d1d5db; box-shadow: 0 4px 12px rgba(0,0,0,0.08); }

    /* ── KPI Tile ── */
    .kpi-tile {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 16px 20px;
        text-align: center;
        position: relative;
        overflow: hidden;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    .kpi-tile:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0,0,0,0.08);
    }
    .kpi-tile::before {
        content: "";
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 4px;
    }
    .kpi-tile.blue::before  { background: #2563eb; }
    .kpi-tile.green::before { background: #10b981; }
    .kpi-tile.amber::before { background: #f59e0b; }
    .kpi-tile.red::before   { background: #ef4444; }
    .kpi-tile.purple::before{ background: #8b5cf6; }
    .kpi-tile.cyan::before  { background: #06b6d4; }
    
    .kpi-label {
        font-size: 0.75rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
        color: #6b7280;
        font-weight: 600;
        margin-bottom: 4px;
    }
    .kpi-value {
        font-size: 1.75rem;
        font-weight: 700;
        color: #111827;
        font-family: 'DM Sans', sans-serif;
        line-height: 1.1;
    }
    .kpi-sub {
        font-size: 0.7rem;
        color: #9ca3af;
        margin-top: 4px;
    }
    .kpi-sub .pos { color: #059669; }
    .kpi-sub .neg { color: #dc2626; }

    /* ── Headers ── */
    .chart-header {
        font-size: 0.75rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        color: #9ca3af;
        margin-bottom: 16px;
        display: flex;
        align-items: center;
        gap: 8px;
    }

    /* ── Record Card ── */
    .entity-card {
        background: #ffffff;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        padding: 16px;
        margin-bottom: 12px;
        border-left: 4px solid #2563eb;
        transition: background 0.2s ease, border-color 0.2s ease;
    }
    .entity-card:hover {
        background: #f9fafb;
    }
    .entity-key { color: #6b7280; font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.02em; }
    .entity-val { color: #111827; font-size: 0.88rem; font-weight: 500; }

    /* ── Sidebar ── */
    section[data-testid="stSidebar"] {
        background: #f7f7f8 !important;
        border-right: 1px solid #e5e7eb;
    }
    
    /* Reduce font size of sidebar elements for a clean, neat appearance */
    section[data-testid="stSidebar"] button p,
    section[data-testid="stSidebar"] button {
        font-size: 0.8rem !important;
    }
    
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3 {
        font-size: 0.95rem !important;
    }
    
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] label p,
    section[data-testid="stSidebar"] div[data-baseweb="select"] {
        font-size: 0.8rem !important;
    }

    section[data-testid="stSidebar"] .stAlert p,
    section[data-testid="stSidebar"] .stAlert span,
    section[data-testid="stSidebar"] .stMarkdown p,
    section[data-testid="stSidebar"] caption,
    section[data-testid="stSidebar"] .stCaptionContainer {
        font-size: 0.8rem !important;
    }
    
    /* ── Chat Messages ── */
    .stChatMessage {
        background: transparent !important;
        border: none !important;
        padding: 24px 0 !important;
    }
    
    /* Assistant Background */
    div[data-testid="stChatMessage"] {
         border-bottom: 1px solid #f0f0f0;
    }

    .stMarkdown p { color: #374151; line-height: 1.75; font-size: 0.95rem; }

    /* ── Streamlit Overrides ── */
    .stAlert { border-radius: 12px !important; border: 1px solid #e5e7eb !important; background: #ffffff !important; }
    .stButton button { border-radius: 8px !important; }
    
    #MainMenu, footer { visibility: hidden; }
    .viewerBadge_container__r5tak { display: none !important; }
    
    header[data-testid="stHeader"] {
        background: #ffffff !important;
        border-bottom: 1px solid #e5e7eb;
    }

    .stTabs [data-baseweb="tab"] {
        font-size: 0.8rem !important;
        font-weight: 600 !important;
        color: #6b7280 !important;
    }
    .stTabs [aria-selected="true"] {
        color: #111827 !important;
        border-bottom-color: #2563eb !important;
    }
</style>
""", unsafe_allow_html=True)

# ─── Session State ─────────────────────────────────────────────────────────────
if "messages" not in st.session_state:
    st.session_state.messages = []
if "mcp_tools" not in st.session_state:
    st.session_state.mcp_tools = get_all_mcp_tools_sync()
if "session_files" not in st.session_state:
    st.session_state.session_files = {}   # {filename: text_content}
if "session_dfs" not in st.session_state:
    st.session_state.session_dfs = {}     # {filename: pd.DataFrame}
if "session_id" not in st.session_state:
    st.session_state.session_id = str(int(time.time()))

# ─── Persistent Chat History Logic ─────────────────────────────────────────────
HISTORY_FILE = "chat_history.json"

def save_chat_session():
    if not st.session_state.messages: return
    try:
        history = get_chat_history()
        # Create a title from the first user message
        first_user_msg = next((m["content"] for m in st.session_state.messages if m["role"] == "user"), "New Chat")
        title = first_user_msg[:40] + ("..." if len(first_user_msg) > 40 else "")
        
        current_sid = st.session_state.session_id
        
        # Check if session exists
        found = False
        for entry in history:
            if entry["id"] == current_sid:
                entry["title"] = title
                entry["messages"] = st.session_state.messages
                found = True
                break
        
        if not found:
            history.insert(0, {
                "id": current_sid, 
                "title": title, 
                "messages": st.session_state.messages,
                "dt": time.ctime()
            })
        
        # Keep only last 25 sessions
        history = history[:25]
        with open(HISTORY_FILE, "w") as f:
            json.dump(history, f)
    except: pass

def get_chat_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r") as f:
                return json.load(f)
        except: return []
    return []

def load_session(session_id):
    history = get_chat_history()
    for entry in history:
        if entry["id"] == session_id:
            st.session_state.messages = entry["messages"]
            st.session_state.session_id = entry["id"]
            return True
    return False

# ─── In-Memory File Text Extraction ────────────────────────────────────────────
FILE_ICONS = {"csv":"📊","xlsx":"📗","xls":"📗","pdf":"📕",
              "docx":"📘","doc":"📘","txt":"📄","md":"📝","json":"🗄️"}

def extract_text(uploaded_file) -> str:
    """Extract plain text from an uploaded file entirely in-memory."""
    name = uploaded_file.name
    ext  = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    raw  = uploaded_file.read()  # bytes

    try:
        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            # Join pages with double newline to preserve paragraph structure
            pages_text = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                # Merge hyphenated line-breaks (e.g. "infor-\nmation" -> "information")
                import re as _pdf_re
                page_text = _pdf_re.sub(r'-(\n)\s*', '', page_text)
                pages_text.append(page_text)
            return "\n\n".join(pages_text), None

        elif ext in ("xlsx", "xls"):
            sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
            out = []
            all_dfs = []
            for sheet, df in sheets.items():
                # SAAS PRODUCTION GRADE: Robust Header Detection
                # Require at least 50% non-empty columns or at least 3 non-empty columns to avoid metadata rows like title/reynolds
                if any("Unnamed" in str(c) for c in df.columns) or df.columns.tolist() == list(range(len(df.columns))):
                    for i in range(min(5, len(df))):
                        row = df.iloc[i]
                        non_empty = row.count()
                        if non_empty >= 2 and non_empty >= min(3, len(row) * 0.4):
                            df.columns = [str(x).strip() for x in row]
                            df = df.iloc[i+1:].reset_index(drop=True)
                            break
                
                # De-duplicate column names to prevent pandas concat InvalidIndexError
                cols = []
                count = {}
                for col in df.columns:
                    c_str = str(col).strip()
                    if not c_str or c_str.lower() == 'nan':
                        c_str = 'Unnamed'
                    if c_str in count:
                        count[c_str] += 1
                        cols.append(f"{c_str}_{count[c_str]}")
                    else:
                        count[c_str] = 0
                        cols.append(c_str)
                df.columns = cols
                
                out.append(f"=== Sheet: {sheet} ===\n{df.to_string(index=False)}")
                all_dfs.append(df)
            
            try:
                combined_df = pd.concat(all_dfs, ignore_index=True) if all_dfs else pd.DataFrame()
            except Exception:
                combined_df = max(all_dfs, key=len) if all_dfs else pd.DataFrame()
            
            return out, combined_df

        elif ext == "csv":
            df = pd.read_csv(io.BytesIO(raw))
            # Similar header logic for CSV
            if any("Unnamed" in str(c) for c in df.columns):
                potential_headers = df.iloc[0]
                if any(isinstance(x, str) and len(x) > 2 for x in potential_headers):
                    df.columns = df.iloc[0]
                    df = df.iloc[1:].reset_index(drop=True)
            
            # De-duplicate CSV columns
            cols = []
            count = {}
            for col in df.columns:
                c_str = str(col).strip()
                if not c_str or c_str.lower() == 'nan':
                    c_str = 'Unnamed'
                if c_str in count:
                    count[c_str] += 1
                    cols.append(f"{c_str}_{count[c_str]}")
                else:
                    count[c_str] = 0
                    cols.append(c_str)
            df.columns = cols
            
            return [df.to_string(index=False)], df

        elif ext in ("docx", "doc"):
            try:
                from docx import Document
                doc = Document(io.BytesIO(raw))
                out = []
                
                # 1. Headers & Footers (Often contains Chapter Titles)
                for section in doc.sections:
                    for header in [section.header, section.first_page_header, section.even_page_header]:
                        if header: out.extend([p.text for p in header.paragraphs if p.text.strip()])
                    for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                        if footer: out.extend([p.text for p in footer.paragraphs if p.text.strip()])

                # 2. Main Body Paragraphs
                out.extend([p.text for p in doc.paragraphs if p.text.strip()])

                # 3. Tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text: out.append(row_text)
                
                # 4. Shapes/Textboxes (Often used for diagrams/notes)
                # Attempt to find text boxes in shapes
                try:
                    for shape in doc.inline_shapes:
                        pass # inline shapes usually don't have text we can jump to easily, but we try
                except: pass

                return "\n\n".join(out), None
            except Exception:
                # .doc (old binary format) — extract printable ASCII as fallback
                import re as _re
                decoded = raw.decode("latin-1", errors="ignore")
                # Extract runs of printable characters (length >= 4)
                words = _re.findall(r'[\x20-\x7E]{4,}', decoded)
                clean = " ".join(words)
                if len(clean) > 100:
                    return clean
                return (
                    "[Warning] Could not fully read this .doc file. "
                    "Please save it as .docx (Word 2007+) and re-upload for best results.\n\n"
                    + clean
                )

        else:  # txt, md, json, etc.
            return raw.decode("utf-8", errors="ignore"), None

    except Exception as e:
        return f"[Error reading {name}: {e}]", None


# ─── In-Memory File Q&A ─────────────────────────────────────────────────────────
def answer_from_session_files(question: str, session_files: dict, llm_model: str) -> dict:
    """Answer a question from in-memory uploaded files. Never reads from disk."""

    if not session_files:
        return {
            "answer": (
                "📁 **No files uploaded yet.**\n\n"
                "Please upload a PDF, Word, Excel, or CSV file using the **Upload Files** panel in the sidebar."
            ),
            "data": None, "sql": "", "visual_hint": {"show": False}
        }

    q_lower = question.lower()
    file_names = list(session_files.keys())

    # ── Pick the most relevant file (or combine all) ───────────────────────────
    # SAAS PRODUCTION GRADE: Advanced Multi-File Relevance Engine (Exact + Overlap + Tabular Schema matching)
    target_name    = None
    target_content = ""
    
    best_fname = None
    best_score = 0
    
    # Pre-clean question words for matching
    import re as _sel_re
    q_words = set(_sel_re.findall(r'[a-z0-9]+', q_lower))
    
    for fname, content in session_files.items():
        fname_l = fname.lower()
        stem = fname.rsplit(".", 1)[0].lower() if "." in fname else fname_l
        
        score = 0
        # 1. Exact matches (highest priority)
        if fname_l in q_lower:
            score += 200
        elif stem in q_lower:
            score += 100
        else:
            # 2. Singular/plural adjustments & file stem token overlap
            stem_tokens = _sel_re.findall(r'[a-z0-9]+', stem)
            for token in stem_tokens:
                # Direct match, or singular/plural variants (e.g. "projects" matches "project")
                if token in q_words or (token.endswith('s') and token[:-1] in q_words) or (token + 's' in q_words):
                    score += 20
            
            # 3. Tabular Schema Overlap: check if question references Excel/CSV column names
            if fname in st.session_state.session_dfs:
                df_cols = st.session_state.session_dfs[fname].columns
                for col in df_cols:
                    col_l = str(col).lower().replace('\n', ' ').strip()
                    col_tokens = _sel_re.findall(r'[a-z0-9]+', col_l)
                    for col_tok in col_tokens:
                        if len(col_tok) > 2 and col_tok in q_words:
                            score += 15  # Column name overlap strongly suggests this Excel/CSV file is the target!
        
        if score > best_score:
            best_score = score
            best_fname = fname

    if best_fname and best_score > 0:
        target_name = best_fname
        target_content = session_files[best_fname]
    else:
        # Fallback: if only one file was uploaded, use it. If multiple and no clear match, combine all.
        if len(session_files) == 1:
            target_name = file_names[0]
            target_content = session_files[target_name]
        else:
            combined = [f"=== {fn} ===\n{ct}" for fn, ct in session_files.items()]
            target_content = "\n\n".join(combined)
            target_name    = ", ".join(file_names)

    # ── Detect meta/overview questions → use document start ───────────────────
    meta_patterns = [
        "about what", "discuss", "about this", "what is this", "what does this",
        "tell me about", "summarize", "summarise", "summary", "overview",
        "introduction", "topic", "subject", "this document", "the document",
        "this report", "the report", "what is it about", "explain this",
        "describe this", "document about", "report about",
    ]
    is_meta = any(p in q_lower for p in meta_patterns)

    # ── Build context — always give the LLM something to work with ─────────────
    # SAAS PRODUCTION GRADE: Pass full context for small/medium documents (<350k chars)
    # This prevents keyword matching failures and gives the LLM 100% accurate global context.
    if len(target_content) < 350000:
        context = target_content
    elif is_meta:
        # Overview: start of document
        paragraphs = [p.strip() for p in target_content.split("\n") if p.strip()]
        context = "\n".join(paragraphs[:50])[:3500]
    else:
        paragraphs = [p.strip() for p in target_content.split("\n") if p.strip()]
        # ── Keyword scoring with Sliding Window ──────────────────────────────────
        import re as regex_match
        
        # Stop-words to ignore when scoring
        stop_words = {"the", "and", "of", "to", "in", "is", "that", "it", "he", "was", "for", "on", "are", "as", "with", "his", "they", "i", "at", "be", "this", "have", "from", "or", "one", "had", "by", "word", "but", "not", "what", "all", "were", "we", "when", "your", "can", "said", "there", "use", "an", "each", "which", "she", "do", "how", "their", "if", "will", "up", "other", "about", "out", "many", "then", "them", "these", "so", "some", "her", "would", "make", "like", "him", "into", "has", "look", "two", "more", "write", "go", "see", "number", "no", "way", "could", "people", "my", "than", "first", "water", "been", "call", "who", "oil", "its", "now", "find", "about", "what", "how", "why", "where", "who", "which"}
        # Extract clean words from question (strip punctuation) for scoring
        import re as _kw_re
        raw_kw_tokens = _kw_re.findall(r'[a-z0-9]+', q_lower)
        keywords = [kw for kw in raw_kw_tokens if len(kw) > 2 and kw not in stop_words]
        
        # Enhanced Chapter Regex (Matches Ch 6, Chapter VI, Section 6.1, etc)
        chapter_pat = r'(chapter|section|part|unit)\s*([0-9]+|[ivxlc]+|[a-z])'
        user_chap = regex_match.search(chapter_pat, q_lower)
        
        scored_indices = []
        import re as _score_re
        for i, p in enumerate(paragraphs):
            score = 0
            # Strip punctuation from paragraph words for fair matching
            p_low_clean = " ".join(_score_re.findall(r'[a-z0-9]+', p.lower()))
            
            if user_chap and user_chap.group(0) in p_low_clean:
                score += 100 # High priority for exact chapter match
            
            # Context-sensitive match: if we found "Chapter 6", prioritize numbers like "6"
            if user_chap and user_chap.group(2) in p_low_clean:
                score += 5
            
            # Score: +3 for exact word match, +1 for partial (substring) match
            for kw in keywords:
                if _score_re.search(r'\b' + _score_re.escape(kw) + r'\b', p_low_clean):
                    score += 3
                elif kw in p_low_clean:
                    score += 1
            
            if score > 0:
                scored_indices.append((score, i))

        if scored_indices:
            # Sort by score, keep top 15 matches
            scored_indices.sort(key=lambda x: x[0], reverse=True)
            top_indices = [idx for score, idx in scored_indices[:15]]
            
            # Build context using Sliding Window (pull in nearby lines for each match)
            final_context_blocks = []
            seen_indices = set()
            
            for idx in top_indices:
                # Take 2 paragraphs before and 6 paragraphs after the match
                for window_idx in range(max(0, idx - 2), min(len(paragraphs), idx + 7)):
                    if window_idx not in seen_indices:
                        final_context_blocks.append(paragraphs[window_idx])
                        seen_indices.add(window_idx)
            
            context = "\n".join(final_context_blocks)[:20000]
        else:
            # Fallback: Just give the beginning (up to 20k chars)
            context = "\n".join(paragraphs[:80])[:20000]

    # ── Structure for LLM ─────────────────────────────────────────────────────
    df_info = ""
    is_tabular = target_name in st.session_state.session_dfs
    if is_tabular:
        df_obj = st.session_state.session_dfs[target_name]
        # Build richer dataframe info: dtypes + sample + numeric cols
        numeric_cols = df_obj.select_dtypes(include='number').columns.tolist()
        df_info = (
            f"DATAFRAME INFO:\n"
            f"- Filename: {target_name}\n"
            f"- Shape: {df_obj.shape[0]} rows x {df_obj.shape[1]} columns\n"
            f"- All Columns: {df_obj.columns.tolist()}\n"
            f"- Numeric Columns (use pd.to_numeric for others): {numeric_cols}\n"
            f"- Data Types:\n{df_obj.dtypes.to_string()}\n"
            f"- First 5 rows:\n{df_obj.head(5).to_string()}\n"
            f"- Last 3 rows:\n{df_obj.tail(3).to_string()}\n"
        )
        prompt = (
            "SYSTEM: You are a Senior SaaS Data Scientist & Python expert. "
            "You MUST write executable Python code to answer the question accurately.\n\n"
            f"FULL DATA CONTEXT:\n{context}\n\n"
            f"DATAFRAME METADATA:\n{df_info}\n\n"
            "CRITICAL CALCULATION RULES (follow ALL of these):\n"
            "1. The dataframe is already loaded as 'df'. Do NOT re-read any file.\n"
            "2. ALWAYS coerce numeric columns before calculating: "
            "   use pd.to_numeric(df['col'], errors='coerce') before .sum()/.mean()/.max()/.min() etc.\n"
            "3. For SUM/TOTAL: use df['col'].pipe(pd.to_numeric, errors='coerce').sum()\n"
            "4. For AVERAGE/MEAN: use df['col'].pipe(pd.to_numeric, errors='coerce').mean()\n"
            "5. For COUNT: use len(df) or df['col'].count() or df['col'].value_counts()\n"
            "6. For PERCENTAGE: compute numerator/denominator*100 with round(val, 2)\n"
            "7. Column names may have spaces, special chars or mixed case — use df.columns to discover them first, "
            "   then access with df['exact_column_name'] or df.iloc[:, index].\n"
            "8. Strip whitespace from string columns before filtering: df['col'].str.strip()\n"
            "9. If filtering rows (e.g. by category/name), use case-insensitive: "
            "   df[df['col'].astype(str).str.strip().str.lower() == 'value']\n"
            "10. ALWAYS set 'result' to the final computed value (a number, string, or DataFrame).\n"
            "11. Write code in ONE ```python block. No external imports beyond pd, np, px, re.\n"
            "12. Optionally set 'insights' (list of strings) and 'fig' (Plotly figure) for richer output.\n"
            "13. If the question cannot be answered from the data, set result = 'NOT_FOUND_IN_DOCUMENTS'.\n\n"
            f"USER QUESTION: {question}\n"
            "PYTHON CODE (executable, sets result=...):"
        )
    else:
        prompt = (
            "You are a document Q&A assistant. Your ONLY job is to answer based on the DOCUMENT CONTEXT below.\n"
            "Rules:\n"
            "1. Read the DOCUMENT CONTEXT carefully.\n"
            "2. Answer the USER QUESTION using information found in the context.\n"
            "3. If the answer is clearly present, give a direct, accurate answer.\n"
            "4. If the answer is NOT in the context at all, reply ONLY with: NOT_FOUND_IN_DOCUMENTS\n"
            "5. Do NOT make up or invent any facts. Do NOT add disclaimers. Just answer.\n"
            "6. Do NOT start your answer with 'ANSWER:', 'Based on the context', or any prefix. Just give the answer directly.\n"
            "7. Smart Entity Mapping: If the user's question uses a slightly imprecise term (such as asking for a 'company' or 'organization' when the document describes 'ports', 'institutions', or related entities), answer with the most matching entity from the document context (e.g. Port of Singapore).\n\n"
            f"DOCUMENT CONTEXT:\n{context}\n\n"
            f"USER QUESTION: {question}\n"
        )

    try:
        raw_answer = _gemini_chat(prompt, llm_model, timeout=300)
        clean = raw_answer.strip()
        # Strip common LLM prefix tags from text answers
        import re as _pfx_re
        clean = _pfx_re.sub(
            r'^(\*{0,2}ANSWER:\*{0,2}|\*{0,2}Answer:\*{0,2}|RESPONSE:|Response:|RESULT:|Result:|OUTPUT:|Output:|Assistant:)\s*',
            '', clean, flags=_pfx_re.IGNORECASE
        ).strip()

        # ── Execution Logic ──
        if "```python" in clean and target_name in st.session_state.session_dfs:
            calc_error = None
            code = ""
            try:
                import re as _re
                import pandas as pd
                import numpy as _np
                import plotly.express as px
                from datetime import datetime as _dt, timedelta as _td
                try:
                    from sklearn.linear_model import LinearRegression
                    from sklearn.cluster import KMeans
                except: pass
                
                # 1. Extraction
                code_match = _re.search(r"```python\n(.*?)```", clean, _re.DOTALL)
                if code_match:
                    code = code_match.group(1)
                    
                    # 2. Aggressive Sanitization
                    lines = code.split("\n")
                    clean_lines = []
                    for line in lines:
                        l_strip = line.strip()
                        if l_strip.startswith(("import pandas", "import sklearn", "df = pd.read", "print(")):
                            continue
                        clean_lines.append(line)
                    code = "\n".join(clean_lines)
                    
                    # 3. Production Helper
                    def parse_duration(val):
                        try:
                            val = str(val).lower().replace(" ", "").replace(":", ".")
                            if "-" not in val: return 0.0
                            parts = val.split("-")
                            s = _dt.strptime(parts[0], "%I.%M%p")
                            e = _dt.strptime(parts[1], "%I.%M%p")
                            if e < s: e += _td(days=1)
                            return (e - s).total_seconds() / 3600.0
                        except: return 0.0

                    # 4. Mock Environment
                    df_live = st.session_state.session_dfs[target_name]
                    exec_scope = {
                        "df": df_live, "pd": pd, "np": _np, "px": px, "re": _re,
                        "datetime": _dt, "timedelta": _td,
                        "parse_duration": parse_duration, 
                        "result": None, "insights": [], "fig": None, "prediction": None
                    }
                    
                    # Try to add sklearn if possible
                    try:
                        from sklearn.linear_model import LinearRegression as LR
                        exec_scope["LinearRegression"] = LR
                    except: pass

                    exec(code, exec_scope)
                    
                    if exec_scope.get("result") is not None:
                        final_res = exec_scope.get("result")
                        # Grounding Check
                        if str(final_res).upper() == "NOT_FOUND_IN_DOCUMENTS":
                            return {
                                "answer": "❌ **Information Not Found in Documents**\n\nI couldn't find any information locally related to your question. Please try asking about something contained in your uploaded files.",
                                "data": None,
                                "sql": f"Checked: {target_name}",
                                "visual_hint": {"show": False}
                            }

                        # Prepare data for return (for persistence in chat history)
                        export_data = None
                        if isinstance(final_res, (pd.DataFrame, pd.Series, list)):
                            export_data = pd.DataFrame(final_res).to_dict(orient='records')
                        
                        ans_text = f"**Analysis Result:** {final_res if not isinstance(final_res, (pd.DataFrame, pd.Series, list)) else 'See analytical breakdown below'}"
                        
                        # Capture AI generated insights, predictions and fig
                        insights = exec_scope.get("insights", [])
                        pred = exec_scope.get("prediction")
                        
                        # ── Serialize LLM-generated Plotly fig to JSON so it can be stored and re-rendered ──
                        fig_json = None
                        llm_fig = exec_scope.get("fig")
                        if llm_fig is not None:
                            try:
                                import plotly.io as _pio
                                _apply_enterprise_figure_theme(llm_fig)
                                fig_json = _pio.to_json(llm_fig)
                            except Exception:
                                fig_json = None
                        
                        return {
                            "answer": ans_text,
                            "data": export_data,
                            "insights": insights,
                            "prediction": pred,
                            "fig_json": fig_json,
                            "sql": f"SaaS Analysis Logic",
                            "visual_hint": {"show": bool(fig_json), "type": "Auto"}
                        }
            except Exception as e:
                calc_error = str(e)
            
            # ── Auto-Retry: ask LLM to fix the code if it errored ─────────────
            if calc_error and code:
                retry_prompt = (
                    f"The following Python code raised an error when run on a pandas DataFrame:\n"
                    f"ERROR: {calc_error}\n\n"
                    f"ORIGINAL CODE:\n```python\n{code}\n```\n\n"
                    f"DATAFRAME INFO:\n{df_info}\n"
                    f"USER QUESTION: {question}\n\n"
                    "Fix the code so it runs correctly. "
                    "Always coerce numeric columns with pd.to_numeric(df['col'], errors='coerce') before aggregating. "
                    "Output ONLY the corrected ```python code block. Set 'result' to the final answer."
                )
                try:
                    retry_raw = _gemini_chat(retry_prompt, llm_model, timeout=120)
                    retry_match = __import__('re').search(r'```python\n(.*?)```', retry_raw, __import__('re').DOTALL)
                    if retry_match:
                        retry_code = retry_match.group(1)
                        retry_scope = {
                            "df": st.session_state.session_dfs[target_name],
                            "pd": pd, "np": _np, "px": px, "re": _re,
                            "datetime": _dt, "timedelta": _td,
                            "parse_duration": parse_duration,
                            "result": None, "insights": [], "fig": None, "prediction": None,
                            "plotly": __import__('plotly')
                        }
                        exec(retry_code, retry_scope)
                        if retry_scope.get("result") is not None:
                            final_res = retry_scope["result"]
                            export_data = None
                            if isinstance(final_res, (pd.DataFrame, pd.Series, list)):
                                export_data = pd.DataFrame(final_res).to_dict(orient='records')
                            ans_text = f"**Analysis Result:** {final_res if not isinstance(final_res, (pd.DataFrame, pd.Series, list)) else 'See analytical breakdown below'}"
                            # Serialize retry fig
                            retry_fig_json = None
                            retry_llm_fig = retry_scope.get("fig")
                            if retry_llm_fig is not None:
                                try:
                                    import plotly.io as _pio2
                                    _apply_enterprise_figure_theme(retry_llm_fig)
                                    retry_fig_json = _pio2.to_json(retry_llm_fig)
                                except Exception:
                                    retry_fig_json = None
                            return {
                                "answer": ans_text,
                                "data": export_data,
                                "insights": retry_scope.get("insights", []),
                                "prediction": retry_scope.get("prediction"),
                                "fig_json": retry_fig_json,
                                "sql": "SaaS Analysis Logic (auto-corrected)",
                                "visual_hint": {"show": bool(retry_fig_json), "type": "Auto"}
                            }
                except Exception:
                    pass
                with st.expander("🔍 Calculation Diagnostic", expanded=False):
                    st.error(f"Error: {calc_error}")
                    st.code(code)

        # Show the actual context use for debugging in the UI
        with st.expander("🔍 View Source Context (Paragraphs Used)", expanded=False):
            st.text(context)

        if "NOT_FOUND" in clean.upper() or len(clean) < 4:
            return {
                "answer": (
                    f"❌ **Not found in `{target_name}`.**\n\n"
                    "The uploaded document does not contain the answer to your question.\n\n"
                    "*Try asking something else, or upload a different file.*"
                ),
                "data": None,
                "sql": f"File: {target_name}",
                "visual_hint": {"show": False}
            }

        return {"answer": clean, "data": None, "sql": f"File: {target_name}", "visual_hint": {"show": False}}

    except Exception as e:
        # LLM failed → return raw content so user still gets something
        fallback = "\n".join(paragraphs[:12])
        return {
            "answer": (
                f"⚠️ *LLM unavailable ({e}). Showing raw document content:*\n\n"
                f"**{target_name}**\n\n{fallback}"
            ),
            "data": None,
            "sql": f"File: {target_name}",
            "visual_hint": {"show": False}
        }

# ─── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    
    # ── New Chat Button ──────────────────────────────────────────────
    if st.button("➕ New Chat", width="stretch"):
        st.session_state.messages = []
        # Reset session ID for a truly new chat
        st.session_state.session_id = str(int(time.time()))
        st.rerun()

    st.divider()

    # ── Chat History List (Persistent Sessions) ──────────────────────────────
    st.markdown("### 🕑 Recent Chats")
    history = get_chat_history()
    
    for entry in history:
        btn_label = entry['title']
        
        if st.button(btn_label, key=f"hist_{entry['id']}", width="stretch"):
            if load_session(entry['id']):
                st.rerun()
            
    if history and st.button("🗑️ Clear All History", width="stretch"):
        if os.path.exists(HISTORY_FILE): os.remove(HISTORY_FILE)
        st.session_state.messages = []
        st.session_state.session_id = str(int(time.time()))
        st.rerun()

    st.divider()

    # ── Gemini Model Selection ─────────────────────────────────────────────────
    from mcp_engine import GEMINI_API_KEY, LLM_MODEL
    model_options = ["gemini-2.5-flash-lite", "gemini-2.5-flash", "gemini-3.5-flash", "gemini-3.1-flash-lite", "gemini-2.5-pro"]
    st.info("🚀 **Gemini API Active**")

    model_index = 0
    if LLM_MODEL in model_options:
        model_index = model_options.index(LLM_MODEL)

    model = st.selectbox(
        "🤖 Gemini Model",
        model_options,
        index=model_index
    )

    # Data Source selection
    data_source = st.selectbox(
        "📊 Data Source", 
        ["Auto", "Postgres", "MySQL", "SharePoint", "Local Files"],
        index=1, # Default to Postgres as requested
        help="Choose a specific system to query, or let the AI detect it automatically."
    )

    # ── File Upload Panel (only when Local Files is chosen) ────────────────────
    if data_source == "Local Files":
        st.divider()
        st.subheader("📎 Upload Files")
        st.caption("Files are read in-memory. Nothing is saved to disk.")

        uploaded = st.file_uploader(
            "Drop files here",
            type=["csv", "xlsx", "xls", "pdf", "docx", "doc", "txt", "md", "json"],
            accept_multiple_files=True,
            label_visibility="collapsed"
        )

        if uploaded:
            new_count = 0
            for f in uploaded:
                if f.name not in st.session_state.session_files:
                    text_data, df_data = extract_text(f)
                    # For sheets list, join if it's a list
                    st.session_state.session_files[f.name] = "\n".join(text_data) if isinstance(text_data, list) else text_data
                    if df_data is not None:
                        st.session_state.session_dfs[f.name] = df_data
                    new_count += 1
            if new_count:
                st.success(f"✅ {new_count} new file(s) loaded into memory")

        # Show currently loaded files
        sf = st.session_state.session_files
        if sf:
            st.markdown(f"**{len(sf)} file(s) loaded this session:**")
            for fname in sf:
                ext  = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
                icon = FILE_ICONS.get(ext, "📎")
                char_count = len(sf[fname])
                st.markdown(f"{icon} `{fname}` — *{char_count:,} chars*")

            if st.button("🗑️ Clear Uploaded Files"):
                st.session_state.session_files = {}
                st.session_state.session_dfs = {}
                st.rerun()
        else:
            st.info("No files loaded. Upload above to get started.")

    st.divider()
    st.subheader("📡 Distributed Core")
    if st.session_state.mcp_tools:
        st.success("✨ **Glimpse AI HTTP Gateway** (SSE/HTTP)")
        st.caption(f"Broadcasting {len(st.session_state.mcp_tools)} tools across company.")
        if st.button("🔄 Refresh Tools", width="stretch"):
            st.session_state.mcp_tools = get_all_mcp_tools_sync()
            st.rerun()
    else:
        st.error("❌ Gateway Offline")
        if st.button("📡 Reconnect to Gateway", width="stretch"):
            st.session_state.mcp_tools = get_all_mcp_tools_sync()
            st.rerun()

    st.divider()


# ─── Premium Header ────────────────────────────────────────────────────────────
top_c1, top_c2 = st.columns([0.7, 0.3])
with top_c1:
    st.title("✨ Glimpse AI")
    st.markdown("##### *Production SaaS Interface — Active Insight Engine*")
with top_c2:
    st.markdown("<br>", unsafe_allow_html=True)
    status_color = "#008a66" if st.session_state.mcp_tools else "#ef4444"
    status_bg = "rgba(0, 204, 150, 0.05)" if st.session_state.mcp_tools else "rgba(239, 68, 68, 0.05)"
    status_border = "rgba(0, 204, 150, 0.4)" if st.session_state.mcp_tools else "rgba(239, 68, 68, 0.4)"
    status_text = "🟢 AGENT ONLINE" if st.session_state.mcp_tools else "🔴 GATEWAY OFFLINE"
    conn_text = "MCP Gateway Connected" if st.session_state.mcp_tools else "Connection Lost"
    
    st.markdown(f"""
        <div style="background: {status_bg}; border: 1px solid {status_border}; border-radius: 12px; padding: 10px; text-align: center;">
            <span style="color: {status_color}; font-weight: 700;">{status_text}</span><br>
            <span style="color: rgba(0,0,0,0.5); font-size: 0.8rem;">{conn_text}</span>
        </div>
    """, unsafe_allow_html=True)

st.write("") # Spacer


# ─── Dashboard helpers ─────────────────────────────────────────────────────────
# ─── Question Intent Classification Engine ──────────────────────────────────
# This classifies every question into a visual mode WITHOUT calling the LLM.
# This is the production-grade approach used by tools like Tableau/PowerBI AI.

def normalize_dashboard_df(df):
    """Convert database numeric strings to chartable metrics without changing labels."""
    if df is None or df.empty:
        return df

    normalized = df.copy()
    identifier_words = ("id", "code", "phone", "mobile", "zip", "pin", "postal")
    for col in normalized.columns:
        if not (pd.api.types.is_object_dtype(normalized[col]) or pd.api.types.is_string_dtype(normalized[col])):
            continue
        col_name = str(col).lower()
        if any(word in col_name for word in identifier_words):
            continue

        original = normalized[col]
        non_null_count = original.notna().sum()
        if non_null_count == 0:
            continue

        cleaned = (
            original.astype("string")
            .str.replace(",", "", regex=False)
            .str.replace("$", "", regex=False)
            .str.strip()
        )
        converted = pd.to_numeric(cleaned, errors="coerce")
        if converted.notna().sum() / non_null_count >= 0.8:
            normalized[col] = converted

    return normalized


def classify_visual_intent(question: str, df) -> dict:
    """
    Returns a dict:
    {
      'mode': 'PLAIN' | 'TABLE' | 'KPI+TABLE' | 'CHART+KPI+TABLE',
      'chart_type': 'Bar' | 'Line' | 'Donut' | 'Area' | 'Treemap' |
                    'Funnel' | 'Scatter' | 'Waterfall' | None,
      'reason': str   # dev debug string
    }
    
    PLAIN       → no data/table shown at all (pure text answer)
    TABLE       → just show the data table (listing, lookup, detail)
    KPI+TABLE   → financial/aggregate queries with totals
    CHART+KPI+TABLE → analytical/comparative/trends (full dashboard)
    """
    df = normalize_dashboard_df(df)
    q = question.lower().strip()
    rows = len(df) if df is not None else 0
    cols = list(df.columns) if df is not None else []

    num_cols = [c for c in cols if pd.api.types.is_numeric_dtype(df[c])] if df is not None else []
    cat_cols = [
        c for c in cols
        if pd.api.types.is_object_dtype(df[c]) or pd.api.types.is_string_dtype(df[c])
    ] if df is not None else []

    # ── 1. PLAIN: conversation / capability questions (no data needed) ────────
    plain_kw = [
        "hello", "hi ", "hey ", "who are you", "what can you", "help me",
        "capabilities", "what data", "how do you"
    ]
    if any(k in q for k in plain_kw):
        return {"mode": "PLAIN", "chart_type": None, "reason": "greeting/meta"}

    # ── 2. PLAIN: single-value aggregate (no visual needed) ───────────────────
    # ── 3. CHART: explicit charting request ───────────────────────────────────
    chart_explicit_kw = [
        "chart", "graph", "plot", "visual", "visualize", "visualise",
        "visualization", "visualisation", "visula", "dashboard",
        "power bi", "powerbi", "show graph", "draw chart",
        "bar chart", "pie chart", "donut", "line chart", "area chart",
        "treemap", "tree map", "funnel", "pipeline", "scatter", "bubble",
        "correlation", "waterfall", "variance", "bridge"
    ]
    if any(k in q for k in chart_explicit_kw):
        chart_type = "Bar"
        if "line" in q or "trend" in q: chart_type = "Line"
        if "area" in q: chart_type = "Area"
        if "pie" in q or "donut" in q: chart_type = "Donut"
        if "treemap" in q or "tree map" in q: chart_type = "Treemap"
        if "funnel" in q or "pipeline" in q: chart_type = "Funnel"
        if any(k in q for k in ["scatter", "bubble", "correlation"]): chart_type = "Scatter"
        if any(k in q for k in ["waterfall", "variance", "bridge"]): chart_type = "Waterfall"
        return {"mode": "CHART+KPI+TABLE", "chart_type": chart_type, "reason": "explicit chart request"}

    # ── 4. CHART: comparison / trend / analysis ───────────────────────────────
    comparison_kw = [
        "profit by", "sales by", "revenue by", "expenses by", "cost by",
        "by region", "by department", "by category", "by product",
        "by customer", "by supplier", "by branch", "by status", "by employee",
        "each department", "per department", "each city", "per city",
        "each customer", "per customer", "each project", "per project",
        "by month", "by year", "by quarter", "by date", "monthly", "quarterly", "yearly",
        "over the years", "over years", "across years",
        "compare", "comparison", "vs ", "versus", "breakdown",
        "trend", "over time", "growth", "decline", "top ", "bottom ", "best ",
        "ranking", "ranked", "distribution", "share"
    ]
    if any(k in q for k in comparison_kw):
        chart_type = "Bar"
        if any(k in q for k in ["trend", "over time", "over the years", "over years",
                                "across years", "monthly", "quarterly", "yearly",
                                "by month", "by year", "by date"]):
            chart_type = "Line"
        if any(k in q for k in ["share", "percent", "proportion", "pie", "donut"]):
            chart_type = "Donut"
        if any(k in q for k in ["pipeline", "funnel", "conversion"]):
            chart_type = "Funnel"
        if any(k in q for k in ["waterfall", "variance", "bridge"]):
            chart_type = "Waterfall"
        return {"mode": "CHART+KPI+TABLE", "chart_type": chart_type, "reason": "comparison/analysis"}

    # Infer charts from natural BI questions when the returned data can support one.
    # Examples: "Which department has more revenue?" or "How are sales changing each month?"
    if num_cols:
        time_words = ["month", "week", "quarter", "year", "date", "day"]
        movement_words = ["change", "changing", "increase", "decrease", "rise", "fall", "performing", "progress"]
        rank_words = [
            "most", "least", "more", "less", "better", "worst", "largest",
            "smallest", "leading", "highest", "lowest", "top", "bottom",
        ]
        comparison_entities = [
            "department", "product", "customer", "region", "branch", "category",
            "supplier", "employee", "team", "location", "city"
        ]

        if rows >= 2 and any(word in q for word in time_words) and any(word in q for word in movement_words):
            return {"mode": "CHART+KPI+TABLE", "chart_type": "Line", "reason": "natural trend question"}

        if (rows >= 2 and any(word in q for word in rank_words) and
                (cat_cols or any(word in q for word in comparison_entities))):
            return {"mode": "CHART+KPI+TABLE", "chart_type": "Bar", "reason": "natural comparison question"}

    single_agg_kw = [
        "total payroll", "total salary", "salary total", "payroll total",
        "total revenue", "total count", "how many", "count of",
        "average", " avg ", "maximum", "minimum", "highest", "lowest",
        "max salary", "min salary", "what is the total",
    ]
    if any(k in q for k in single_agg_kw):
        if rows <= 1 or q.startswith(("what is", "what's", "how many", "which is")):
            return {"mode": "PLAIN", "chart_type": None, "reason": "single-value aggregate"}
        return {"mode": "KPI+TABLE", "chart_type": None, "reason": "ranked/single metric result"}

    # ── 5. CHART: if data has both numeric+categorical cols and 2+ rows ───────
    # ── 6. KPI+TABLE: queries with numeric data (aggregate-ish) ──────────────
    kpi_kw = [
        "payroll", "salary", "wage", "income", "inventory", "stock",
        "invoice", "payment", "budget", "revenue", "sales", "profit",
        "expense", "cost", "price", "amount"
    ]
    if num_cols and any(k in q for k in kpi_kw):
        return {"mode": "KPI+TABLE", "chart_type": None, "reason": "numeric/financial data"}

    # ── 7. TABLE: listing / lookup (the most common query type) ──────────────
    # "show all employees", "list products", "get customers", etc.
    listing_kw = [
        "show", "list", "get", "display", "fetch", "view",
        "all employees", "all products", "all customers", "all users",
        "all orders", "all invoices", "all projects", "all departments",
        "all payroll", "all sales", "all inventory", "all suppliers",
        "employees", "members", "staff", "personnel",
    ]
    # ── 7. TABLE with Smart Visuals: listing with grouping potential ────────
    if any(k in q for k in listing_kw):
        return {"mode": "TABLE", "chart_type": None, "reason": "pure listing"}

    # ── 8. If we have data with numeric cols, show KPI+TABLE ──────────────────
    if num_cols:
        return {"mode": "KPI+TABLE", "chart_type": None, "reason": "has metrics"}

    # ── 9. Default: just TABLE ────────────────────────────────────────────────
    return {"mode": "TABLE", "chart_type": None, "reason": "default"}


def resolve_visual_intent(question: str, df, visual_hint=None) -> dict:
    """Use one grounded decision for chart display after data retrieval."""
    intent = classify_visual_intent(question, df)
    if not isinstance(visual_hint, dict) or "show" not in visual_hint:
        return intent

    df = normalize_dashboard_df(df)
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist() if df is not None else []
    if not visual_hint.get("show"):
        if intent["mode"] == "CHART+KPI+TABLE":
            mode = "KPI+TABLE" if num_cols else "TABLE"
            return {"mode": mode, "chart_type": None, "reason": "grounded response does not require chart"}
        return intent

    hint_type = str(visual_hint.get("type", "")).title()
    valid_chart_types = {"Bar", "Line", "Pie", "Donut", "Area", "Treemap",
                         "Funnel", "Scatter", "Waterfall"}
    if hint_type not in valid_chart_types:
        return intent

    if not num_cols:
        return {"mode": "TABLE", "chart_type": None, "reason": "no numerical chart metric"}

    supported, reason = _supports_meaningful_chart(df, hint_type, visual_hint)
    if not supported:
        return {"mode": "KPI+TABLE", "chart_type": None, "reason": reason}

    return {"mode": "CHART+KPI+TABLE", "chart_type": hint_type, "reason": "grounded visual decision"}


def _supports_meaningful_chart(df, chart_type, visual_hint):
    """Require enough comparable values for an interpretable business visual."""
    metric = visual_hint.get("y")
    dimension = visual_hint.get("x")
    if metric not in df.columns or not pd.api.types.is_numeric_dtype(df[metric]):
        return False, "visual metric is unavailable"

    if chart_type in {"Line", "Area"}:
        if dimension not in df.columns or df[dimension].dropna().astype(str).nunique() < 2:
            return False, "time series needs multiple periods"
        return True, None

    if chart_type == "Scatter":
        numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_columns) < 2 or len(df) < 3:
            return False, "scatter chart needs multiple numeric observations"
        return True, None

    if dimension not in df.columns or df[dimension].dropna().astype(str).nunique() < 2:
        return False, "comparison chart needs multiple groups"
    if df[metric].dropna().nunique() < 2:
        return False, "comparison values are tied"
    return True, None


def get_meaningful_cols(df):
    """Filter and rank columns by business relevance."""
    junk = ["id", "pk", "index", "phone", "mobile", "zip", "pin", "postal",
            "serial", "fax", "created_at", "updated_at", "version"]
    
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    cat_cols = df.select_dtypes(include=["object"]).columns.tolist()
    
    # Priority-based metric scoring
    metric_priority = ["salary", "amount", "revenue", "profit", "cost", "total", "pay", "price", "budget", "count"]
    def metric_score(c):
        c_low = c.lower()
        if any(x in c_low for x in junk): return -100
        for i, p in enumerate(metric_priority):
            if p in c_low: return len(metric_priority) - i
        return 0
    
    clean_num = sorted(
        [c for c in num_cols if metric_score(c) >= 0],
        key=metric_score,
        reverse=True
    )
    
    # Priority-based dimension scoring
    dim_priority = ["name", "title", "role", "dept", "department", "category", "status", "type", "region", "city"]
    def dim_score(c):
        c_low = c.lower()
        if any(x in c_low for x in junk): return -100
        for i, p in enumerate(dim_priority):
            if p in c_low: return len(dim_priority) - i
        return 0

    clean_cat = sorted(
        [c for c in cat_cols if dim_score(c) >= 0],
        key=dim_score,
        reverse=True
    )
    
    return clean_num, clean_cat


def _plotly_layout():
    return dict(
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='#ffffff',
        font=dict(family="Inter", size=11, color="#4b5563"),
        title_font=dict(size=14, family="Inter", color="#111827"),
        margin=dict(t=26, b=38, l=14, r=14),
        colorway=BI_COLORS,
        hoverlabel=dict(bgcolor="#ffffff", font=dict(family="Inter", size=12, color="#111827")),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="left", x=0,
                    bgcolor='rgba(0,0,0,0)', borderwidth=0),
        xaxis=dict(gridcolor="#eef2f7", linecolor="#d1d5db", tickfont=dict(color="#6b7280")),
        yaxis=dict(gridcolor="#eef2f7", linecolor="#d1d5db", tickfont=dict(color="#6b7280")),
    )


def resolve_chart_title(title, x_col, y_col):
    """Return a readable title when generated chart metadata is empty or invalid."""
    normalized = str(title or "").strip().strip("\"'").lower()
    invalid_titles = {"", "undefined", "null", "none", "n/a", "na", "data analysis", "chart", "graph"}
    if normalized in invalid_titles:
        y_label = str(y_col).replace('_', ' ').title()
        x_label = str(x_col).replace('_', ' ').title()
        return f"{y_label} by {x_label}"
    return str(title).strip().strip("\"'")


BI_COLORS = [
    "#118DFF", "#12239E", "#E66C37", "#6B007B", "#E044A7",
    "#744EC2", "#D9B300", "#D64550", "#197278", "#1AAB40"
]
TILE_CLASSES = ["blue", "green", "amber", "purple", "red", "cyan"]


def _apply_enterprise_figure_theme(fig):
    """Apply a consistent Power BI-like palette to generated Plotly figures."""
    fig.update_layout(**_plotly_layout(), coloraxis_showscale=False, height=420)
    for index, trace in enumerate(fig.data):
        color = BI_COLORS[index % len(BI_COLORS)]
        if trace.type == "pie":
            trace.update(marker=dict(colors=BI_COLORS, line=dict(color="#ffffff", width=2)))
        elif trace.type in {"bar", "funnel"}:
            if len(fig.data) == 1:
                count = len(trace.x) if trace.x is not None else len(trace.y)
                trace.update(marker=dict(color=[BI_COLORS[i % len(BI_COLORS)] for i in range(count)]))
            else:
                trace.update(marker=dict(color=color))
        elif trace.type in {"scatter", "scattergl"}:
            trace.update(marker=dict(color=color), line=dict(color=color))
    return fig


def render_kpi_tiles(df):
    """Power BI-style KPI metric tiles — only for numeric data."""
    num_cols, cat_cols = get_meaningful_cols(df)
    if not num_cols:
        # Fallback for listings: show record count as a KPI
        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
        c1, _ = st.columns([1, 4])
        with c1:
            st.markdown(f"""
                <div class="kpi-tile blue">
                    <div class="kpi-label">Total Records</div>
                    <div class="kpi-value">{len(df)}</div>
                    <div class="kpi-sub">entries analyzed</div>
                </div>
            """, unsafe_allow_html=True)
        st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)
        return
    st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
    n = min(len(num_cols), 5)
    kpi_cols = st.columns(n)
    for i, col in enumerate(num_cols[:n]):
        val = df[col].sum()
        avg = df[col].mean()
        tile_cls = TILE_CLASSES[i % len(TILE_CLASSES)]
        disp = f"{val:,.0f}" if val >= 1 else f"{val:.4f}"
        kpi_cols[i].markdown(f"""
            <div class="kpi-tile {tile_cls}">
                <div class="kpi-label">{col.replace('_',' ')}</div>
                <div class="kpi-value">{disp}</div>
                <div class="kpi-sub">avg {avg:,.1f}</div>
            </div>
        """, unsafe_allow_html=True)
    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)


def render_data_table(df):
    """Always-visible styled data table."""
    if df.empty:
        return
    num_cols, _ = get_meaningful_cols(df)
    try:
        styled = df.style.background_gradient(cmap="Blues", subset=num_cols[:2] if num_cols else [])
        st.dataframe(styled, width="stretch", height=min(400, 38 + len(df) * 36))
    except Exception:
        st.dataframe(df, width="stretch")


def render_chart(df, question, chart_type="Bar", visual_hint=None):
    """Render an enterprise analytical visual with robust axis resolution and grouping."""
    if df.empty:
        return
    num_cols, cat_cols = get_meaningful_cols(df)
    
    if not num_cols and not cat_cols:
        return  # Purely textual data

    layout = _plotly_layout()

    # Resolve X axis (category) & Y axis (metric) with strictly prioritized business logic
    visual_hint = visual_hint or {}
    x_hint = visual_hint.get("x")
    y_hint = visual_hint.get("y")
    q_low = question.lower()
    
    x_col, y_col = None, None
    col_map_low = {c.lower(): c for c in df.columns}

    # Step 1: Metric Selection (Y-Axis)
    # 1.1 Match AI hint (exact or substring)
    if y_hint and y_hint != "Count":
        y_hint_l = y_hint.lower()
        if y_hint_l in col_map_low:
            y_col = col_map_low[y_hint_l]
        else:
            for col_l, col in col_map_low.items():
                if y_hint_l in col_l or col_l in y_hint_l:
                    y_col = col
                    break

    # 1.2 Search for numeric columns mentioned in the question
    if not y_col:
        for col in num_cols:
            col_l = col.lower()
            if col_l in q_low or col_l.replace('_', ' ') in q_low:
                y_col = col
                break

    # 1.3 Only calculate frequencies when the returned dataset does not
    # already contain the requested grouped count metric.
    has_metric_hint = y_col is not None
    is_count_query = (y_hint == "Count") or (
        not has_metric_hint and any(k in q_low for k in ["how many", "count", "number of"])
    )

    # 1.4 Fallback to highest ranked numeric column (only if not a count query)
    if not y_col and not is_count_query and num_cols:
        y_col = num_cols[0]

    # Step 2: Dimension Selection (X-Axis)
    # 2.1 Try AI hint first (exact or substring)
    if x_hint:
        x_hint_l = x_hint.lower()
        if x_hint_l in col_map_low:
            x_col = col_map_low[x_hint_l]
        else:
            for col_l, col in col_map_low.items():
                if x_hint_l in col_l or col_l in x_hint_l:
                    x_col = col
                    break

    # 2.2 Date columns for time-series charts
    if not x_col and (chart_type in {"Line", "Area"} or "trend" in q_low or "over time" in q_low):
        date_cols = [c for c in df.columns if any(x in c.lower() for x in ["date", "month", "year", "time", "day"])]
        if date_cols:
            x_col = date_cols[0]

    # 2.3 Search for dimensions in the question
    if not x_col:
        for col in cat_cols:
            col_l = col.lower()
            if col_l in q_low or col_l.replace('_', ' ') in q_low:
                x_col = col
                break

    # 2.4 Fallback to highest ranked category
    if not x_col and cat_cols:
        x_col = cat_cols[0]

    # Scatter plots need two numerical axes; prefer a second available metric.
    scatter_x_col = next((col for col in num_cols if col != y_col), None)
    if chart_type == "Scatter" and scatter_x_col:
        x_col = scatter_x_col

    # ── Final aggregation & fallback setup ──
    if not y_col or y_col == "Count" or is_count_query:
        # Group by the resolved X-axis and count frequencies
        y_col = "Count"
        if not x_col:
            x_col = cat_cols[0] if cat_cols else df.columns[0]
        chart_df = df.groupby(x_col).size().reset_index(name="Count")
    else:
        # Standard metric mapping
        if not x_col:
            non_num_cols = [c for c in df.columns if c != y_col and c not in num_cols]
            x_col = non_num_cols[0] if non_num_cols else (df.columns[0] if df.columns[0] != y_col else df.columns[-1])
        chart_df = df.copy()

    # Guard: x and y must not be the same column
    if x_col == y_col:
        if len(df.columns) > 1:
            x_col = [c for c in df.columns if c != y_col][0]
        elif y_col in num_cols:
            # Explicit visual requests for a scalar metric still need an axis label.
            x_col = "Metric"
            chart_df = df.copy()
            chart_df[x_col] = y_col.replace('_', ' ').title()
        else:
            return  # Cannot draw a meaningful chart

    # Chart title - generated title metadata may be empty or an invalid placeholder.
    title_text = resolve_chart_title(visual_hint.get("title"), x_col, y_col)
    st.markdown(f"""<p class="chart-header">📊 &nbsp;{title_text}</p>""", unsafe_allow_html=True)

    if not 'chart_df' in locals():
        chart_df = df.copy()
    # Sort for better readability
    try:
        ascending = visual_hint.get("performance_direction") == "lowest"
        chart_df = chart_df.sort_values(y_col, ascending=ascending)
    except Exception:
        pass

    try:
        if chart_type in {"Line", "Area"}:
            try: chart_df = chart_df.sort_values(x_col)
            except: pass
            if chart_type == "Area":
                fig = px.area(chart_df.head(30), x=x_col, y=y_col,
                              color_discrete_sequence=["#118DFF"],
                              labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})
                fig.update_traces(line=dict(width=2.5, color="#118DFF"),
                                  fillcolor="rgba(17,141,255,0.18)")
            else:
                fig = px.line(chart_df.head(30), x=x_col, y=y_col, markers=True,
                              color_discrete_sequence=["#197278"],
                              labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})
                fig.update_traces(line=dict(width=3), marker=dict(size=7, color="#197278"))

        elif chart_type in {"Pie", "Donut"}:
            top_df = chart_df.nlargest(10, y_col)
            fig = px.pie(top_df, values=y_col, names=x_col, hole=0.45,
                         color_discrete_sequence=BI_COLORS)
            fig.update_traces(textposition='inside', textinfo='percent+label',
                              marker=dict(line=dict(color="#ffffff", width=2)))

        elif chart_type == "Treemap":
            top_df = chart_df.nlargest(20, y_col)
            fig = px.treemap(top_df, path=[x_col], values=y_col, color=x_col,
                             color_discrete_sequence=BI_COLORS)
            fig.update_traces(textinfo="label+value+percent parent",
                              marker=dict(line=dict(color="#ffffff", width=2)))

        elif chart_type == "Funnel":
            funnel_df = chart_df.nlargest(12, y_col)
            fig = px.funnel(funnel_df, y=x_col, x=y_col, color=x_col,
                            color_discrete_sequence=BI_COLORS,
                            labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})
            fig.update_traces(textposition="inside", textinfo="value+percent initial")

        elif chart_type == "Scatter" and scatter_x_col:
            size_col = next((col for col in num_cols if col not in {scatter_x_col, y_col}), None)
            fig = px.scatter(chart_df.head(100), x=scatter_x_col, y=y_col,
                             color=x_col if x_col in cat_cols else None,
                             size=size_col,
                             color_discrete_sequence=BI_COLORS,
                             labels={scatter_x_col: scatter_x_col.replace('_', ' '),
                                     y_col: y_col.replace('_', ' ')})
            fig.update_traces(marker=dict(opacity=0.82, line=dict(width=1, color="#ffffff")))

        elif chart_type == "Waterfall":
            waterfall_df = chart_df.head(15)
            fig = go.Figure(go.Waterfall(
                x=waterfall_df[x_col].astype(str),
                y=waterfall_df[y_col],
                measure=["relative"] * len(waterfall_df),
                connector={"line": {"color": "#cbd5e1"}},
                increasing={"marker": {"color": "#1AAB40"}},
                decreasing={"marker": {"color": "#D64550"}},
                totals={"marker": {"color": "#118DFF"}},
                textposition="outside"
            ))

        else:  # Default: Bar
            # Production SaaS: Always limit to top 15 for readability, unless explicitly scrolled
            plot_df = chart_df.head(15) if len(chart_df) > 15 else chart_df
            
            # Use horizontal bars if too many items or long labels
            orientation = 'v'
            if len(plot_df) > 8 or plot_df[x_col].astype(str).str.len().max() > 12:
                orientation = 'h'
                # Swap X and Y for horizontal
                fig = px.bar(plot_df, x=y_col, y=x_col, text_auto='.3s',
                             color=x_col, orientation='h',
                             color_discrete_sequence=BI_COLORS,
                             labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})
                fig.update_layout(yaxis={'categoryorder':'total ascending'})
            else:
                fig = px.bar(plot_df, x=x_col, y=y_col, text_auto='.3s',
                             color=x_col,
                             color_discrete_sequence=BI_COLORS,
                             labels={x_col: x_col.replace('_', ' '), y_col: y_col.replace('_', ' ')})

        # Suppress internal Plotly title to avoid "undefined" and use our styled header instead
        fig.update_layout(**layout, title=None, coloraxis_showscale=False, height=420)
        st.plotly_chart(fig, width="stretch", config={'displayModeBar': False})

    except Exception as e:
        st.caption(f"Chart could not be rendered: {e}")


def _format_analysis_value(column, value):
    """Format a calculated value for a grounded business insight."""
    if pd.isna(value):
        return "N/A"
    if isinstance(value, (int, float, np.integer, np.floating)):
        financial = ("amount", "salary", "budget", "sales", "revenue", "cost", "price", "pay", "profit")
        if any(word in str(column).lower() for word in financial):
            return f"${float(value):,.2f}"
        if float(value).is_integer():
            return f"{int(value):,}"
        return f"{float(value):,.2f}"
    return str(value)


def build_grounded_analysis(df, chart_type, visual_hint=None):
    """Produce insight text and a cautious projection from displayed rows."""
    if df is None or df.empty:
        return None, None

    visual_hint = visual_hint or {}
    num_cols, cat_cols = get_meaningful_cols(df)
    y_col = visual_hint.get("y") if visual_hint.get("y") in df.columns else (num_cols[0] if num_cols else None)
    x_col = visual_hint.get("x") if visual_hint.get("x") in df.columns else (cat_cols[0] if cat_cols else None)
    if not y_col or not pd.api.types.is_numeric_dtype(df[y_col]):
        return None, None

    values = df.dropna(subset=[y_col]).copy()
    if values.empty:
        return None, None

    y_label = str(y_col).replace("_", " ")
    if chart_type in {"Line", "Area"} and len(values) >= 2:
        start = float(values[y_col].iloc[0])
        end = float(values[y_col].iloc[-1])
        delta = end - start
        direction = "increased" if delta > 0 else "decreased" if delta < 0 else "remained unchanged"
        period_text = f" from {values[x_col].iloc[0]} to {values[x_col].iloc[-1]}" if x_col else ""
        percentage = f" ({abs(delta / start) * 100:,.1f}%)" if start and delta else ""
        insight = (
            f"{y_label.title()} {direction}{period_text}: "
            f"{_format_analysis_value(y_col, start)} to {_format_analysis_value(y_col, end)}"
        )
        if delta:
            insight += f", a change of {_format_analysis_value(y_col, abs(delta))}{percentage}."
        else:
            insight += "."

        note = visual_hint.get("measure_note")
        if note:
            insight = f"{note} {insight}"

        forecast = None
        if len(values) >= 3 and visual_hint.get("forecast_periods"):
            average_change = (end - start) / (len(values) - 1)
            horizon = int(visual_hint.get("forecast_periods", 1) or 1)
            projected_values = [max(0.0, end + average_change * step) for step in range(1, horizon + 1)]
            projected = sum(projected_values) if horizon > 1 else projected_values[0]
            label = visual_hint.get("forecast_label", "next period")
            forecast = (
                f"Simple {label} projection: {_format_analysis_value(y_col, projected)} "
                f"based on average change across {len(values) - 1} observed intervals."
            )
        return insight, forecast

    if x_col and len(values) >= 2:
        is_lowest = visual_hint.get("performance_direction") == "lowest"
        ranked = values.sort_values(y_col, ascending=is_lowest)
        focus = ranked.iloc[0]
        comparison = ranked.iloc[1]
        gap = abs(float(focus[y_col]) - float(comparison[y_col]))
        if is_lowest:
            insight = (
                f"{focus[x_col]} is the lowest performer in {y_label} at {_format_analysis_value(y_col, focus[y_col])}, "
                f"trailing the next-lowest group by {_format_analysis_value(y_col, gap)}."
            )
        else:
            insight = (
                f"{focus[x_col]} leads in {y_label} at {_format_analysis_value(y_col, focus[y_col])}, "
                f"ahead of {comparison[x_col]} by {_format_analysis_value(y_col, gap)}."
            )
        if visual_hint.get("dimension_note"):
            insight = f"{visual_hint['dimension_note']} {insight}"
        return insight, None

    return f"{y_label.title()}: {_format_analysis_value(y_col, values[y_col].iloc[0])}.", None



def render_response(df, question, model, visual_hint=None, cached_insights=None, cached_forecast=None):
    """
    Master render function — decides exactly what to show based on question intent.
    Returns (insights, forecast) for caching.
    """
    if df is None or df.empty:
        return None, None

    df = normalize_dashboard_df(df)
    intent = resolve_visual_intent(question, df, visual_hint)
    mode = intent["mode"]
    chart_type = intent["chart_type"]

    insights = cached_insights
    forecast = cached_forecast

    if mode == "PLAIN":
        return None, None

    # ── MODE: TABLE only (listings, lookups, directory queries) ──────────────
    if mode == "TABLE":
        num_cols, _ = get_meaningful_cols(df)
        # Only show record count tile — no KPIs, no charts
        st.markdown(f"""
            <div style="display:flex; gap:12px; margin-bottom:12px;">
                <div style="background:#f0f9ff; border:1px solid #bae6fd; border-radius:8px;
                            padding:8px 16px; font-size:0.82rem; color:#0369a1; font-weight:600;">
                    📋 {len(df)} records found
                </div>
                <div style="background:#f0fdf4; border:1px solid #bbf7d0; border-radius:8px;
                            padding:8px 16px; font-size:0.82rem; color:#166534; font-weight:600;">
                    🗂️ {len(df.columns)} columns
                </div>
            </div>
        """, unsafe_allow_html=True)
        render_data_table(df)
        return None, None

    # ── MODE: KPI + TABLE (financial, payroll, inventory) ────────────────────
    elif mode == "KPI+TABLE":
        render_kpi_tiles(df)
        render_data_table(df)
        return None, None

    # ── MODE: CHART + KPI + TABLE (comparative, trend, analytical) ───────────
    elif mode == "CHART+KPI+TABLE":
        # 1. KPI row
        render_kpi_tiles(df)
        # 2. Chart
        render_chart(df, question, chart_type, visual_hint)
        # 3. Interpret the returned rows without depending on LLM availability.
        if not insights:
            insights, grounded_forecast = build_grounded_analysis(df, chart_type, visual_hint)
            if not forecast:
                forecast = grounded_forecast


        if insights:
            safe_insights = html.escape(str(insights)).replace("\n", "<br>")
            safe_forecast = html.escape(str(forecast)).replace("\n", "<br>") if forecast else None
            c1, c2 = st.columns([0.6, 0.4])
            with c1:
                st.markdown(f"""
                    <div style="background:rgba(37,99,235,0.04); border-left:4px solid #2563eb;
                                padding:14px; border-radius:8px; border:1px solid rgba(37,99,235,0.12);">
                        <div style="font-weight:700; color:#1e3a8a; margin-bottom:8px; font-size:0.78rem;
                                    text-transform:uppercase; letter-spacing:0.06em;">🧠 Agent Insights</div>
                        <div style="font-size:0.88rem; color:#374151; line-height:1.6;">{safe_insights}</div>
                    </div>
                """, unsafe_allow_html=True)
            if forecast:
                with c2:
                    st.markdown(f"""
                        <div style="background:rgba(16,185,129,0.04); border-left:4px solid #10b981;
                                    padding:14px; border-radius:8px; border:1px solid rgba(16,185,129,0.12);">
                            <div style="font-weight:700; color:#065f46; margin-bottom:8px; font-size:0.78rem;
                                        text-transform:uppercase; letter-spacing:0.06em;">🔮 Forecast</div>
                            <div style="font-size:0.88rem; color:#065f46; line-height:1.6; font-style:italic;">{safe_forecast}</div>
                        </div>
                    """, unsafe_allow_html=True)

        # 4. Data table (collapsed by default when chart is shown)
        with st.expander("🗄️ View Full Data Table", expanded=False):
            render_data_table(df)
        return insights, forecast

    # Fallback — show table
    render_data_table(df)
    return None, None


# ─── Chat Interface ─────────────────────────────────────────────────────────────
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        if m["role"] == "assistant":
            content = m['content']
            # Production Polish: Filter out common LLM artifacts or empty responses
            clean = content.replace("**ANSWER:** ", "").replace("ANSWER: ", "").strip()
            if clean and clean.lower() != "undefined":
                st.markdown(clean)
        else:
            st.markdown(m["content"])
        sql_val = m.get("sql", "")
        if sql_val and sql_val not in ("General Conversation", "", None):
            is_tool = any(str(sql_val).startswith(p) for p in ("Tool:", "File:", "File "))
            label   = "🔧 Source" if is_tool else "🔍 SQL Query"
            with st.expander(label, expanded=False):
                st.code(sql_val, language="text" if is_tool else "sql")
        # Replay generated figures only when the original question merits a chart.
        display_fig_json = m.get("fig_json")
        if display_fig_json and m.get("data"):
            orig_q = m.get("orig_question", m.get("content", ""))
            saved_df = normalize_dashboard_df(pd.DataFrame(m["data"]))
            if resolve_visual_intent(orig_q, saved_df, m.get("visual_hint"))["mode"] != "CHART+KPI+TABLE":
                display_fig_json = None
        if display_fig_json:
            try:
                import plotly.io as _pio_hist
                _hist_fig = _pio_hist.from_json(display_fig_json)
                orig_q = m.get("orig_question", m.get("content", ""))
                _auto_title = orig_q[:80].strip().rstrip('?').title()
                _apply_enterprise_figure_theme(_hist_fig)
                _hist_fig.update_layout(title=dict(text=_auto_title, font=dict(size=13, family='Inter', color='#111827'), x=0))
                st.plotly_chart(_hist_fig, width="stretch",
                                config={'displayModeBar': False},
                                key=f"histfig_{m.get('id', id(m))}")
            except Exception:
                pass
        if m.get("data") and len(m["data"]) > 0:
            try:
                _df = normalize_dashboard_df(pd.DataFrame(m["data"]))
                orig_q = m.get("orig_question", m["content"])
                if not display_fig_json:
                    render_response(
                        _df, orig_q, model,
                        visual_hint=m.get("visual_hint"),
                        cached_insights=m.get("insights"),
                        cached_forecast=m.get("prediction")
                    )
                else:
                    render_kpi_tiles(_df)
                    with st.expander("🗄️ View Full Data Table", expanded=False):
                        render_data_table(_df)
                st.download_button(
                    label="📥 Export CSV",
                    data=_df.to_csv(index=False).encode('utf-8'),
                    file_name=f"glimpse_export_{int(time.time())}.csv",
                    mime='text/csv',
                    key=f"dl_{m.get('id', hash(m['content']))}"
                )
            except Exception as _e:
                st.dataframe(pd.DataFrame(m["data"]), width="stretch")


prompt = st.chat_input("Ask a question about your data or uploaded files...")

if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        spinner_msg = (
            "📄 Searching uploaded files..."
            if data_source == "Local Files"
            else f"🔍 Analysing {data_source} data..."
            if data_source != "Auto"
            else "🤖 Consulting Intelligence..."
        )
        with st.spinner(spinner_msg):
            try:
                # ── Local Files: answer from in-memory content ─────────────
                if data_source == "Local Files":
                    res = answer_from_session_files(
                        prompt,
                        st.session_state.session_files,
                        model
                    )
                # ── All other sources: go through MCP engine ───────────────
                else:
                    res = mcp_answer(
                        prompt,
                        st.session_state.mcp_tools,
                        model,
                        data_source=data_source
                    )

                ans      = str(res.get("answer", "No answer returned.") or "No content found.")
                data     = res.get("data", None)
                sql      = res.get("sql", "")
                v_hint   = res.get("visual_hint", {})
                fig_json = res.get("fig_json", None)
                if fig_json and data and len(data) > 0:
                    result_df = normalize_dashboard_df(pd.DataFrame(data))
                    if resolve_visual_intent(prompt, result_df, v_hint)["mode"] != "CHART+KPI+TABLE":
                        fig_json = None

                # ── 1. Primary Answer Text ────────────────────────────────────
                # Render answer only if it's meaningful
                if ans and ans.lower() != "undefined" and len(ans) > 2:
                    st.markdown(ans)

                # ── 2. SQL Source (collapsed) ─────────────────────────────────
                if sql and sql not in ("General Conversation", "", None):
                    is_tool = any(str(sql).startswith(p) for p in ("Tool:", "File:", "File "))
                    label   = "🔧 Source" if is_tool else "🔍 SQL Query"
                    with st.expander(label, expanded=False):
                        st.code(sql, language="text" if is_tool else "sql")

                # ── 3. Data Visualization (intent-driven) ─────────────────────
                saved_insights, saved_forecast = None, None

                # ── 3a. Use LLM-generated Plotly fig directly if available ────
                if fig_json:
                    try:
                        import plotly.io as _pio_render
                        _fig = _pio_render.from_json(fig_json)
                        # Derive a clean title from the question
                        _auto_title = prompt[:80].strip().rstrip('?').title()
                        _apply_enterprise_figure_theme(_fig)
                        _fig.update_layout(title=dict(text=_auto_title, font=dict(size=13, family='Inter', color='#111827'), x=0)
                        )
                        st.plotly_chart(_fig, width="stretch", config={'displayModeBar': False})
                    except Exception as _fig_err:
                        st.caption(f"Chart render error: {_fig_err}")

                # ── 3b. Fallback: build chart from result DataFrame ────────────
                if data and len(data) > 0:
                    df = normalize_dashboard_df(pd.DataFrame(data))
                    try:
                        # Only call render_response if LLM fig was not shown
                        if not fig_json:
                            saved_insights, saved_forecast = render_response(
                                df, prompt, model,
                                visual_hint=v_hint
                            )
                        else:
                            # Show KPI tiles + table even when fig_json is present
                            render_kpi_tiles(df)
                            with st.expander("🗄️ View Full Data Table", expanded=False):
                                render_data_table(df)
                        st.download_button(
                            label="📥 Export CSV",
                            data=df.to_csv(index=False).encode('utf-8'),
                            file_name=f"export_{int(time.time())}.csv",
                            mime='text/csv'
                        )
                    except Exception as e:
                        st.warning(f"⚠️ Visualization error: {e}")
                        st.dataframe(df, width="stretch")

                st.session_state.messages.append({
                    "role": "assistant",
                    "content": ans,
                    "data": data,
                    "orig_question": prompt,
                    "insights": saved_insights,
                    "prediction": saved_forecast,
                    "visual_hint": v_hint,
                    "fig_json": fig_json,
                    "sql": sql
                })

            except Exception as e:
                err_msg = f"❌ Analysis failed: {e}"
                st.error(err_msg)
                st.session_state.messages.append({"role": "assistant", "content": err_msg})
        # Save to history file for persistence
        save_chat_session()
        st.rerun()
